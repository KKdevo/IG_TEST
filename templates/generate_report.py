#!/usr/bin/env python3
"""
Social Media Report Generator
Reads a Word document template and generates a beautiful HTML report.

Usage:
    python generate_report.py input.docx output.html
    python generate_report.py input.docx  (outputs to input_report.html)
"""

import sys
import os
import re
from docx import Document
from docx.table import Table
from datetime import datetime, timedelta
import calendar

def extract_hyperlinks_from_cell(cell, doc_rels):
    """Extract all hyperlink URLs from a cell using doc relationships AND field codes."""
    import re
    hyperlinks = []

    try:
        # Get cell XML
        cell_xml = cell._tc.xml

        # Method 1: Find all r:id references in the cell XML (standard hyperlinks)
        rids = re.findall(r'r:id="(rId\d+)"', cell_xml)

        # Look up each rId in the document relationships
        for rid in rids:
            if rid in doc_rels:
                rel = doc_rels[rid]
                # Check if it's a hyperlink
                if 'hyperlink' in rel.reltype.lower():
                    hyperlinks.append(rel.target_ref)
        
        # Method 2: Extract field code hyperlinks (HYPERLINK "url" in instrText)
        # This handles links pasted from browser or created differently in Word
        # Handle both straight quotes (") and curly quotes (" ")
        field_links = re.findall(r'HYPERLINK\s+["\u201c]([^"\u201d]+)["\u201d]', cell_xml)
        for link in field_links:
            # Unescape XML entities
            link = link.replace('&amp;', '&')
            if link not in hyperlinks:
                hyperlinks.append(link)
                
    except:
        pass

    return hyperlinks


def extract_tables(doc):
    """Extract all tables from the document."""
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)
    return tables


def extract_tables_with_hyperlinks(doc):
    """Extract all tables from the document, including embedded hyperlinks."""
    # Get document relationships for hyperlink lookup
    doc_rels = doc.part.rels
    
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                text = cell.text.strip()
                hyperlinks = extract_hyperlinks_from_cell(cell, doc_rels)
                # If we have hyperlinks, use them instead of/alongside text
                if hyperlinks:
                    # Join all hyperlinks with newlines for URLs
                    cells.append('\n'.join(hyperlinks))
                else:
                    cells.append(text)
            rows.append(cells)
        tables.append(rows)
    return tables

def parse_config(tables):
    """Parse the config table (first table with Account Name/Date Created/Month)."""
    config = {"AccountName": "", "DateCreated": "", "Month": ""}
    for table in tables:
        for row in table:
            if len(row) >= 2:
                key = row[0].strip().lower().replace(" ", "")
                value = row[1].strip()
                if "accountname" in key:
                    config["AccountName"] = value
                elif "datecreated" in key:
                    config["DateCreated"] = value
                elif "month" in key or "reviewmonth" in key:
                    config["Month"] = value
    # Clean placeholder text
    for k, v in config.items():
        if v.startswith("[") and v.endswith("]"):
            config[k] = ""
    return config


def parse_date(date_str, year=None):
    """
    Parse date string in various formats to datetime object.
    Supports: MM/DD, MM/DD/YY, MM/DD/YYYY, Dec 5, December 5, etc.
    """
    if not date_str:
        return None
    
    date_str = date_str.strip()
    
    # Try MM/DD/YY or MM/DD/YYYY format first
    for fmt in ["%m/%d/%y", "%m/%d/%Y"]:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    
    # Try MM/DD format (without year) - common in schedules
    if re.match(r'^\d{1,2}/\d{1,2}$', date_str):
        try:
            parsed = datetime.strptime(date_str, "%m/%d")
            if year:
                parsed = parsed.replace(year=year)
            else:
                parsed = parsed.replace(year=datetime.now().year)
            return parsed
        except ValueError:
            pass
    
    # Try "Dec 5" or "December 5" format
    month_formats = [
        ("%b %d", True),   # Dec 5
        ("%B %d", True),   # December 5
        ("%b %d, %Y", False),  # Dec 5, 2025
        ("%B %d, %Y", False),  # December 5, 2025
    ]
    
    for fmt, needs_year in month_formats:
        try:
            parsed = datetime.strptime(date_str, fmt)
            if needs_year and year:
                parsed = parsed.replace(year=year)
            elif needs_year:
                parsed = parsed.replace(year=datetime.now().year)
            return parsed
        except ValueError:
            continue
    
    return None


def parse_month_year(month_str):
    """
    Parse month string like 'December 2025' or 'Dec 2025' to (month_num, year).
    Returns (month_number, year) tuple or (current_month, current_year) if parsing fails.
    """
    if not month_str:
        now = datetime.now()
        return (now.month, now.year)
    
    month_str = month_str.strip()
    
    # Try "December 2025" or "Dec 2025" format
    for fmt in ["%B %Y", "%b %Y"]:
        try:
            parsed = datetime.strptime(month_str, fmt)
            return (parsed.month, parsed.year)
        except ValueError:
            continue
    
    # Try "12/2025" or "12-2025" format
    for sep in ['/', '-']:
        if sep in month_str:
            parts = month_str.split(sep)
            if len(parts) == 2:
                try:
                    month = int(parts[0])
                    year = int(parts[1])
                    if 1 <= month <= 12:
                        return (month, year)
                except ValueError:
                    continue
    
    # Default to current month/year
    now = datetime.now()
    return (now.month, now.year)


def group_posts_by_week(posts, month, year):
    """
    Group posts by week number within the month.
    Returns dict: {week_num: [(day, post), ...]}
    """
    # Get calendar info for the month
    cal = calendar.Calendar(firstweekday=6)  # Sunday start
    month_days = list(cal.itermonthdays(year, month))
    
    # Calculate week boundaries
    weeks = {}
    current_week = 1
    week_posts = []
    
    for i, day in enumerate(month_days):
        if day == 0:
            continue
        if i > 0 and i % 7 == 0:
            if week_posts:
                weeks[current_week] = week_posts
            current_week += 1
            week_posts = []
    
    # Now assign posts to weeks
    weeks = {}
    for post in posts:
        post_date = parse_date(post.get("PostDate", ""), year)
        if post_date and post_date.month == month and post_date.year == year:
            day = post_date.day
            # Determine which week this day falls in
            week_num = (day - 1) // 7 + 1
            if week_num not in weeks:
                weeks[week_num] = []
            weeks[week_num].append((day, post))
    
    # Sort posts within each week by day
    for week_num in weeks:
        weeks[week_num].sort(key=lambda x: x[0])
    
    return weeks


def get_week_date_range(week_num, month, year):
    """
    Get the date range string for a week (e.g., 'Dec 1-7').
    """
    # Calculate start and end days for this week
    start_day = (week_num - 1) * 7 + 1
    end_day = min(start_day + 6, calendar.monthrange(year, month)[1])
    
    month_abbr = calendar.month_abbr[month]
    return f"{month_abbr} {start_day}-{end_day}"

def parse_posts_table(table):
    """Parse the posting schedule table by reading headers dynamically."""
    posts = []
    if not table or len(table) < 2:
        return posts
    
    # Build header index map (case-insensitive)
    header_row = table[0]
    header_map = {}
    for idx, h in enumerate(header_row):
        h_lower = h.lower().strip()
        if "title" in h_lower:
            header_map["title"] = idx
        elif "date" in h_lower:
            header_map["date"] = idx
        elif "time" in h_lower:
            header_map["time"] = idx
        elif "type" in h_lower:
            header_map["type"] = idx
        elif "status" in h_lower:
            header_map["status"] = idx
        elif "hashtag" in h_lower:
            header_map["hashtags"] = idx
        elif "note" in h_lower:
            header_map["notes"] = idx
        elif "caption" in h_lower:
            header_map["caption"] = idx
        elif "photo" in h_lower:
            # "Photo Links" column for media URLs
            header_map["photo_links"] = idx
        elif "link" in h_lower:
            header_map["link"] = idx
    
    # Parse rows
    for i, row in enumerate(table):
        if i == 0:  # Skip header row
            continue
        
        # Get title - skip if empty or placeholder
        title_idx = header_map.get("title", 0)
        title = row[title_idx].strip() if title_idx < len(row) else ""
        if not title or title.startswith("["):
            continue
        
        post = {"Title": title}
        
        if "date" in header_map and header_map["date"] < len(row):
            post["PostDate"] = row[header_map["date"]].strip()
        if "time" in header_map and header_map["time"] < len(row):
            post["Time"] = row[header_map["time"]].strip()
        if "type" in header_map and header_map["type"] < len(row):
            post["Type"] = row[header_map["type"]].strip().lower().replace("[", "").replace("]", "")
        else:
            post["Type"] = "post"
        if "status" in header_map and header_map["status"] < len(row):
            post["Status"] = row[header_map["status"]].strip().replace("[", "").replace("]", "")
        else:
            post["Status"] = "Draft"
        if "hashtags" in header_map and header_map["hashtags"] < len(row):
            post["Hashtags"] = row[header_map["hashtags"]].strip()
        if "notes" in header_map and header_map["notes"] < len(row):
            post["Notes"] = row[header_map["notes"]].strip()
        if "caption" in header_map and header_map["caption"] < len(row):
            caption = row[header_map["caption"]].strip()
            if caption and not caption.startswith("["):
                post["Caption"] = caption
        if "photo_links" in header_map and header_map["photo_links"] < len(row):
            photo_links = row[header_map["photo_links"]].strip()
            if photo_links and not photo_links.startswith("["):
                post["MediaURL"] = photo_links
        if "link" in header_map and header_map["link"] < len(row):
            link = row[header_map["link"]].strip()
            # Only use Link column for MediaURL if Photo Links wasn't already set
            if link and not link.startswith("[") and not post.get("MediaURL"):
                post["MediaURL"] = link
        
        posts.append(post)
    return posts

def parse_post_blocks(tables):
    """Parse individual post detail blocks (Title/Caption/Hashtags/URL format)."""
    posts = []
    current_type = "post"
    
    for table in tables:
        # Check if this is a post detail block (has Title, Caption rows)
        row_labels = [row[0].strip().lower() if len(row) >= 2 else "" for row in table]
        
        if "title" in row_labels and ("caption" in row_labels or "description" in row_labels):
            post = {"Type": current_type, "Caption": "", "Hashtags": "", "MediaURL": ""}
            
            for row in table:
                if len(row) >= 2:
                    label = row[0].strip().lower()
                    value = row[1].strip()
                    
                    # Skip placeholder text
                    if value.startswith("[") and value.endswith("]"):
                        continue
                        
                    if label == "title":
                        post["Title"] = value
                    elif label in ["caption", "description"]:
                        post["Caption"] = value
                    elif label == "hashtags":
                        post["Hashtags"] = value
                    elif label in ["image url", "video url", "cover image url", "media url"]:
                        post["MediaURL"] = value
            
            if post.get("Title"):
                posts.append(post)
    
    return posts

def parse_stories_table(table):
    """Parse the stories table by reading headers dynamically."""
    stories = []
    if not table or len(table) < 2:
        return stories
    
    # Build header index map
    header_row = table[0]
    header_map = {}
    for idx, h in enumerate(header_row):
        h_lower = h.lower().strip()
        if "title" in h_lower:
            header_map["title"] = idx
        elif "date" in h_lower:
            header_map["date"] = idx
        elif "time" in h_lower:
            header_map["time"] = idx
        elif "interaction" in h_lower:
            header_map["interaction"] = idx
        elif "note" in h_lower:
            header_map["notes"] = idx
        elif "link" in h_lower:
            header_map["link"] = idx
    
    # Parse rows
    for i, row in enumerate(table):
        if i == 0:
            continue
        
        title_idx = header_map.get("title", 0)
        title = row[title_idx].strip() if title_idx < len(row) else ""
        if not title or title.startswith("["):
            continue
        
        story = {"Title": title}
        
        if "date" in header_map and header_map["date"] < len(row):
            story["PostDate"] = row[header_map["date"]].strip()
        if "time" in header_map and header_map["time"] < len(row):
            story["Time"] = row[header_map["time"]].strip()
        if "interaction" in header_map and header_map["interaction"] < len(row):
            story["InteractiveElements"] = row[header_map["interaction"]].strip()
        if "notes" in header_map and header_map["notes"] < len(row):
            story["Notes"] = row[header_map["notes"]].strip()
        if "link" in header_map and header_map["link"] < len(row):
            link = row[header_map["link"]].strip()
            if link and not link.startswith("["):
                story["MediaURL"] = link
        
        stories.append(story)
    return stories

def parse_interactions_table(table):
    """Parse the interactions table."""
    interactions = []
    
    for i, row in enumerate(table):
        if i == 0:  # Skip header
            continue
        if len(row) >= 11 and row[0] and not row[0].startswith("["):
            interaction = {
                "AccountName": row[0].replace("@", "").replace("[", "").replace("]", ""),
                "Platform": row[1].replace("[", "").replace("]", ""),
                "InteractionType": row[2].replace("[", "").replace("]", ""),
                "DailyGoal": row[3].replace("[", "").replace("]", ""),
                "Mon": "TRUE" if row[4].strip().lower() in ["x", "✓", "✔", "true", "yes"] else "FALSE",
                "Tue": "TRUE" if row[5].strip().lower() in ["x", "✓", "✔", "true", "yes"] else "FALSE",
                "Wed": "TRUE" if row[6].strip().lower() in ["x", "✓", "✔", "true", "yes"] else "FALSE",
                "Thu": "TRUE" if row[7].strip().lower() in ["x", "✓", "✔", "true", "yes"] else "FALSE",
                "Fri": "TRUE" if row[8].strip().lower() in ["x", "✓", "✔", "true", "yes"] else "FALSE",
                "Sat": "TRUE" if row[9].strip().lower() in ["x", "✓", "✔", "true", "yes"] else "FALSE",
                "Sun": "TRUE" if row[10].strip().lower() in ["x", "✓", "✔", "true", "yes"] else "FALSE",
            }
            interactions.append(interaction)
    return interactions

def identify_tables(tables, doc):
    """Identify which table is which based on content/headers."""
    result = {
        "config": None,
        "posts_schedule": None,
        "post_blocks": [],
        "stories": None,
        "interactions": None
    }
    
    for i, table in enumerate(tables):
        if not table or not table[0]:
            continue
            
        first_row = [c.lower().strip() for c in table[0]]
        
        # Config table (2-column with Account Name / Date Created in first column)
        if len(table[0]) == 2 and len(table) <= 5:
            labels = [row[0].lower().strip() if len(row) >= 2 else "" for row in table]
            if "account name" in labels or "date created" in labels:
                result["config"] = table
                continue
        
        # Posts schedule (has Title and Type columns - main posting schedule)
        # This is typically the first large table with 7 columns
        if "title" in first_row and "type" in first_row and len(table) > 5:
            if result["posts_schedule"] is None:
                result["posts_schedule"] = table
                continue
        
        # Stories table (has Title, Post Date, Interaction but no Type column)
        # Usually the second large schedule table
        if "title" in first_row and any("interaction" in c for c in first_row) and "type" not in first_row:
            result["stories"] = table
            continue
            
        # Interactions table (has Account, Daily Goal, Mon/Tue/etc columns)
        if any("account" in c for c in first_row) and any("goal" in c for c in first_row) and any("mon" in c for c in first_row):
            result["interactions"] = table
            continue
            
        # Post detail block (Title/Caption/Hashtags format - 2 column small tables)
        if len(table) >= 3 and len(table[0]) == 2:
            labels = [row[0].lower().strip() if len(row) >= 2 else "" for row in table]
            if "title" in labels and ("caption" in labels or "description" in labels):
                # Skip placeholder templates
                title_value = ""
                for row in table:
                    if len(row) >= 2 and row[0].lower().strip() == "title":
                        title_value = row[1].strip()
                if title_value and not title_value.startswith("["):
                    result["post_blocks"].append(table)
    
    return result

def detect_post_type_from_context(tables, block_index):
    """Try to determine post type based on surrounding content."""
    # This is a simple heuristic - in practice you might need to parse paragraph text
    # For now, we'll look at the block's URL field
    if block_index < len(tables):
        table = tables[block_index]
        for row in table:
            if len(row) >= 2:
                label = row[0].lower().strip()
                value = row[1].lower()
                if "video" in label or "youtube" in value or "vimeo" in value:
                    return "reel"
                if "cover" in label:
                    return "highlight"
    return "post"

def generate_html(config, posts, stories, interactions):
    """Generate the final HTML report."""
    
    # Count posts by type
    posts_count = len([p for p in posts if p.get("Type", "").lower() == "post"])
    reels_count = len([p for p in posts if p.get("Type", "").lower() == "reel"])
    highlights_count = len([p for p in posts if p.get("Type", "").lower() == "highlight"])
    
    # Parse month for calendar
    month, year = parse_month_year(config.get("Month", ""))
    
    # Generate calendar views (include both posts AND stories)
    monthly_calendar_html = render_monthly_calendar(posts, stories, month, year)
    weekly_view_html = render_weekly_view(posts, stories, month, year)
    
    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Social Media Content Schedule - {config.get("AccountName", "Report")}</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Bebas+Neue&family=Montserrat:wght@900&display=swap" rel="stylesheet">
    <style>
        :root {{
            --bg-primary: #FAF9F7;
            --bg-secondary: #FFFFFF;
            --text-primary: #1A1A1A;
            --text-secondary: #6B6B6B;
            --text-muted: #9A9A9A;
            --accent: #2D2D2D;
            --accent-warm: #C4A484;
            --border: #E8E6E3;
            --border-light: #F0EFED;
            --shadow: rgba(0,0,0,0.04);
            --shadow-md: rgba(0,0,0,0.08);
        }}
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: 'Manrope', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; background: var(--bg-primary); color: var(--text-primary); line-height: 1.6; font-size: 15px; -webkit-font-smoothing: antialiased; }}
        
        .nav {{ position: fixed; top: 0; left: 0; width: 240px; height: 100vh; background: var(--bg-secondary); border-right: 1px solid var(--border); padding: 40px 24px; overflow-y: auto; z-index: 100; transition: transform 0.3s ease; }}
        .nav-logo {{ width: 100%; max-width: 160px; margin: -55px auto -45px auto; display: block; }}
        .nav-logo img {{ width: 100%; height: auto; display: block; }}
        
        /* Title Page Brand Image */
        .title-brand-image {{ position: absolute; right: -50px; top: calc(50% + 40px); transform: translateY(-50%); height: 172%; max-height: 1210px; width: auto; opacity: 0.9; pointer-events: none; z-index: 0; }}
        @media (max-width: 1024px) {{ .title-brand-image {{ right: -100px; height: 121%; max-height: 807px; }} }}
        @media (max-width: 640px) {{ .title-brand-image {{ right: -175px; height: 91%; max-height: 564px; opacity: 0.7; }} }}
        
        /* Hamburger Menu */
        .hamburger {{ display: none; position: fixed; top: 20px; right: 20px; z-index: 200; width: 40px; height: 40px; background: rgba(255,255,255,0.95); border: none; border-radius: 10px; cursor: pointer; flex-direction: column; align-items: center; justify-content: center; gap: 5px; box-shadow: 0 2px 12px rgba(0,0,0,0.15); backdrop-filter: blur(10px); }}
        .hamburger .bar {{ display: block; width: 18px; height: 2px; background: var(--text-primary); border-radius: 2px; transition: all 0.3s ease; }}
        .hamburger.active .bar:nth-child(1) {{ transform: rotate(45deg) translate(5px, 5px); }}
        .hamburger.active .bar:nth-child(2) {{ opacity: 0; }}
        .hamburger.active .bar:nth-child(3) {{ transform: rotate(-45deg) translate(5px, -5px); }}
        .nav-overlay {{ display: none; position: fixed; inset: 0; background: rgba(0,0,0,0.5); z-index: 90; opacity: 0; transition: opacity 0.3s ease; pointer-events: none; }}
        .nav-overlay.active {{ display: block; opacity: 1; pointer-events: auto; }}
        .nav-subtitle {{ font-size: 11px; text-transform: uppercase; letter-spacing: 1.5px; color: var(--text-muted); margin-bottom: 24px; text-align: center; margin-top: 0; }}
        .nav-section:first-of-type {{ margin-top: 42px; }}
        .nav-section {{ font-family: 'Bebas Neue', sans-serif; font-size: 14px; text-transform: uppercase; letter-spacing: 2px; color: var(--text-muted); margin: 32px 0 12px; }}
        .nav a {{ display: block; color: var(--text-secondary); text-decoration: none; padding: 10px 0; font-size: 14px; transition: all 0.2s ease; }}
        .nav a:hover {{ color: var(--text-primary); }}
        .nav .sub-link {{ padding-left: 16px; font-size: 13px; color: var(--text-muted); }}
        
        .main {{ margin-left: 240px; min-height: 100vh; }}
        section {{ scroll-margin-top: 20px; }}
        
        .title-page {{ min-height: 100vh; display: flex; flex-direction: column; position: relative; overflow: hidden; background: linear-gradient(165deg, #1A1A1A 0%, #2D2D2D 40%, #3D3D3D 100%); }}
        .title-page::before {{ content: ''; position: absolute; inset: 0; background-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 400 400' xmlns='http://www.w3.org/2000/svg'%3E%3Cfilter id='noiseFilter'%3E%3CfeTurbulence type='fractalNoise' baseFrequency='0.9' numOctaves='4' stitchTiles='stitch'/%3E%3C/filter%3E%3Crect width='100%25' height='100%25' filter='url(%23noiseFilter)'/%3E%3C/svg%3E"); opacity: 0.03; pointer-events: none; }}
        .title-header {{ padding: 48px 64px; display: flex; justify-content: space-between; align-items: flex-start; position: relative; z-index: 1; }}
        .title-badge {{ display: inline-flex; align-items: center; gap: 8px; padding: 10px 18px; background: rgba(255,255,255,0.08); border: 1px solid rgba(255,255,255,0.12); border-radius: 100px; backdrop-filter: blur(10px); }}
        .title-badge-dot {{ width: 6px; height: 6px; background: #7FE5A2; border-radius: 50%; animation: pulse 2s ease-in-out infinite; }}
        @keyframes pulse {{ 0%, 100% {{ opacity: 1; transform: scale(1); }} 50% {{ opacity: 0.6; transform: scale(0.9); }} }}
        .title-badge span {{ font-size: 11px; text-transform: uppercase; letter-spacing: 1.5px; color: rgba(255,255,255,0.7); font-weight: 500; }}
        .title-date {{ font-size: 13px; color: rgba(255,255,255,0.5); }}
        .title-content {{ flex: 1; display: flex; flex-direction: column; justify-content: center; padding: 0 64px; position: relative; z-index: 1; }}
        .title-eyebrow {{ font-family: 'Bebas Neue', sans-serif; font-size: 16px; text-transform: uppercase; letter-spacing: 3px; color: var(--accent-warm); margin-bottom: 24px; }}
        .title-main {{ font-family: 'Montserrat', sans-serif; font-size: clamp(48px, 8vw, 64px); font-weight: 900; color: #FFFFFF; line-height: 1.1; letter-spacing: 0; margin-bottom: 24px; max-width: 800px; text-transform: uppercase; }}
        .title-description {{ font-size: 17px; color: rgba(255,255,255,0.5); max-width: 500px; line-height: 1.7; font-weight: 300; }}
        .title-footer {{ padding: 48px 64px; display: flex; justify-content: space-between; align-items: flex-end; position: relative; z-index: 1; }}
        .title-meta {{ display: flex; gap: 64px; }}
        .title-meta-item {{ position: relative; }}
        .title-meta-label {{ font-size: 10px; text-transform: uppercase; letter-spacing: 2px; color: rgba(255,255,255,0.35); margin-bottom: 8px; font-weight: 500; }}
        .title-meta-value {{ font-size: 18px; color: #FFFFFF; font-weight: 500; }}
        .title-decoration {{ position: absolute; right: 64px; bottom: 50%; transform: translateY(50%); width: 320px; height: 320px; border: 1px solid rgba(255,255,255,0.08); border-radius: 50%; pointer-events: none; }}
        .title-decoration::before {{ content: ''; position: absolute; top: 40px; left: 40px; right: 40px; bottom: 40px; border: 1px solid rgba(255,255,255,0.05); border-radius: 50%; }}
        
        .content-section {{ padding: 80px 64px; border-bottom: 1px solid var(--border); }}
        .section-header {{ margin-bottom: 48px; max-width: 600px; }}
        .section-number {{ font-family: 'Bebas Neue', sans-serif; font-size: 18px; text-transform: uppercase; letter-spacing: 3px; color: var(--text-muted); margin-bottom: 16px; }}
        .section-title {{ font-family: 'Montserrat', sans-serif; font-size: 32px; font-weight: 900; letter-spacing: 0; margin-bottom: 12px; line-height: 1.2; display: inline-flex; align-items: center; gap: 12px; text-transform: uppercase; }}
        .section-desc {{ font-size: 15px; color: var(--text-secondary); line-height: 1.7; }}
        
        .table-wrapper {{ background: var(--bg-secondary); border-radius: 16px; overflow: hidden; box-shadow: 0 1px 3px var(--shadow), 0 4px 20px var(--shadow); margin-bottom: 48px; overflow-x: auto; -webkit-overflow-scrolling: touch; }}
        .table-wrapper::-webkit-scrollbar {{ height: 8px; }}
        .table-wrapper::-webkit-scrollbar-track {{ background: var(--bg-primary); border-radius: 4px; }}
        .table-wrapper::-webkit-scrollbar-thumb {{ background: var(--border); border-radius: 4px; }}
        .table-wrapper::-webkit-scrollbar-thumb:hover {{ background: var(--text-muted); }}
        .table-wrapper table {{ min-width: 700px; }}
        table {{ width: 100%; border-collapse: collapse; }}
        th {{ background: var(--bg-primary); padding: 12px 16px; text-align: left; font-size: 10px; text-transform: uppercase; letter-spacing: 1.5px; color: var(--text-muted); font-weight: 600; border-bottom: 1px solid var(--border); }}
        td {{ padding: 12px 16px; border-bottom: 1px solid var(--border-light); font-size: 14px; vertical-align: middle; }}
        tr:last-child td {{ border-bottom: none; }}
        tr:hover {{ background: var(--bg-primary); }}
        
        .status {{ display: inline-flex; align-items: center; gap: 6px; padding: 6px 12px; border-radius: 100px; font-size: 12px; font-weight: 600; }}
        .status::before {{ content: ''; width: 6px; height: 6px; border-radius: 50%; }}
        .status-draft {{ background: #FEF9E7; color: #9A7B2A; }}
        .status-draft::before {{ background: #F4C430; }}
        .status-approved {{ background: #E8F5E9; color: #2E7D32; }}
        .status-approved::before {{ background: #4CAF50; }}
        .status-posted {{ background: #E3F2FD; color: #1565C0; }}
        .status-posted::before {{ background: #2196F3; }}
        .status-needsrevision {{ background: #FFEBEE; color: #C62828; }}
        .status-needsrevision::before {{ background: #EF5350; }}
        
        .type-badge {{ display: inline-block; padding: 5px 12px; border-radius: 6px; font-size: 11px; text-transform: uppercase; letter-spacing: 0.5px; font-weight: 600; }}
        .type-post {{ background: #F3F4F6; color: #374151; }}
        .type-reel {{ background: #FDF2F8; color: #BE185D; }}
        .type-highlight {{ background: #FEF3C7; color: #B45309; }}
        .type-story {{ background: #EDE9FE; color: #7C3AED; }}
        
        .posts-grid {{ display: grid; grid-template-columns: repeat(auto-fill, minmax(min(340px, 100%), 1fr)); gap: 32px; margin-bottom: 48px; }}
        .post-card {{ background: var(--bg-secondary); border-radius: 20px; overflow: hidden; box-shadow: 0 1px 3px var(--shadow), 0 8px 32px var(--shadow); transition: all 0.3s ease; }}
        .post-card:hover {{ transform: translateY(-4px); box-shadow: 0 4px 12px var(--shadow-md), 0 16px 48px var(--shadow-md); }}
        .post-card-media {{ width: 100%; aspect-ratio: 1; object-fit: cover; background: var(--border-light); }}
        .post-card-media.video-container {{ position: relative; background: #000; }}
        .post-card-media iframe {{ width: 100%; height: 100%; border: none; }}
        .post-card-body {{ padding: 24px; }}
        .post-card-header {{ display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 16px; gap: 12px; }}
        .post-card-title {{ font-size: 17px; font-weight: 600; color: var(--text-primary); line-height: 1.3; }}
        .post-card-date {{ font-size: 12px; color: var(--text-muted); margin-bottom: 12px; font-weight: 500; }}
        .post-card-caption {{ font-size: 14px; color: var(--text-secondary); margin-bottom: 16px; display: -webkit-box; -webkit-line-clamp: 4; -webkit-box-orient: vertical; overflow: hidden; line-height: 1.6; }}
        .post-card-hashtags {{ font-size: 13px; color: var(--accent-warm); font-weight: 500; }}
        .no-media {{ width: 100%; aspect-ratio: 1; background: linear-gradient(135deg, var(--border-light) 0%, var(--border) 100%); display: flex; align-items: center; justify-content: center; color: var(--text-muted); font-size: 13px; }}
        
        /* Video Placeholder Styles */
        .video-placeholder {{ width: 100%; aspect-ratio: 1; background: linear-gradient(135deg, #1A1A1A 0%, #2D2D2D 50%, #3D3D3D 100%); display: flex; align-items: center; justify-content: center; cursor: pointer; transition: all 0.3s ease; position: relative; overflow: hidden; }}
        .video-placeholder::before {{ content: ''; position: absolute; inset: 0; background: radial-gradient(circle at center, rgba(255,255,255,0.05) 0%, transparent 70%); pointer-events: none; }}
        .video-placeholder:hover {{ background: linear-gradient(135deg, #2D2D2D 0%, #3D3D3D 50%, #4D4D4D 100%); transform: scale(1.02); }}
        .video-placeholder-content {{ display: flex; flex-direction: column; align-items: center; gap: 16px; color: #fff; text-align: center; padding: 24px; z-index: 1; }}
        .video-placeholder-content svg {{ color: var(--accent-warm); opacity: 0.9; transition: transform 0.3s ease; }}
        .video-placeholder:hover .video-placeholder-content svg {{ transform: scale(1.1); }}
        .video-placeholder-text {{ font-size: 14px; font-weight: 600; opacity: 0.9; }}
        .video-placeholder-filename {{ font-size: 11px; color: rgba(255,255,255,0.5); max-width: 200px; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }}
        
        /* Video Player Styles - Instagram Style */
        .video-player-container {{ position: relative; width: 100%; aspect-ratio: 1; background: #000; overflow: hidden; cursor: pointer; }}
        .video-player {{ width: 100%; height: 100%; object-fit: cover; background: #000; }}
        .video-fallback {{ display: none; position: absolute; inset: 0; background: linear-gradient(135deg, #1A1A1A 0%, #2D2D2D 50%, #3D3D3D 100%); align-items: center; justify-content: center; cursor: pointer; }}
        .video-player-container:has(.video-player[error]) .video-fallback {{ display: flex; }}
        
        /* Play/Pause Indicator */
        .video-play-indicator {{ position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); width: 70px; height: 70px; background: rgba(0,0,0,0.5); border-radius: 50%; display: flex; align-items: center; justify-content: center; opacity: 1; transition: opacity 0.3s ease, transform 0.2s ease; pointer-events: none; }}
        .video-play-indicator svg {{ width: 28px; height: 28px; margin-left: 4px; }}
        .video-player-container.playing .video-play-indicator {{ opacity: 0; transform: translate(-50%, -50%) scale(0.8); }}
        .video-player-container.playing:hover .video-play-indicator {{ opacity: 0; }}
        .video-player-container.show-pause .video-play-indicator {{ opacity: 1; transform: translate(-50%, -50%) scale(1); }}
        .video-player-container.show-pause .video-play-indicator svg {{ margin-left: 0; }}
        
        /* Instagram-style Mute Button */
        .video-mute-btn {{ position: absolute; bottom: 16px; right: 16px; width: 32px; height: 32px; background: rgba(0,0,0,0.6); border: none; border-radius: 50%; display: flex; align-items: center; justify-content: center; cursor: pointer; transition: all 0.2s ease; z-index: 10; }}
        .video-mute-btn:hover {{ background: rgba(0,0,0,0.8); transform: scale(1.1); }}
        .video-mute-btn svg {{ width: 16px; height: 16px; stroke: white; }}
        .video-mute-btn .unmute-icon {{ display: none; }}
        .video-mute-btn .mute-icon {{ display: block; }}
        .video-player-container.unmuted .video-mute-btn .unmute-icon {{ display: block; }}
        .video-player-container.unmuted .video-mute-btn .mute-icon {{ display: none; }}
        
        /* Carousel Styles */
        .carousel {{ position: relative; width: 100%; aspect-ratio: 1; background: var(--border-light); overflow: hidden; }}
        .carousel-container {{ display: flex; width: 100%; height: 100%; transition: transform 0.3s ease; }}
        .carousel-slide {{ flex-shrink: 0; width: 100%; height: 100%; }}
        .carousel-slide img {{ width: 100%; height: 100%; object-fit: cover; max-width: 100%; }}
        .carousel-btn {{ position: absolute; top: 50%; transform: translateY(-50%); width: 36px; height: 36px; border-radius: 50%; background: rgba(255,255,255,0.9); border: none; font-size: 20px; cursor: pointer; display: flex; align-items: center; justify-content: center; opacity: 0; transition: opacity 0.2s ease; z-index: 10; box-shadow: 0 2px 8px rgba(0,0,0,0.15); }}
        .carousel:hover .carousel-btn {{ opacity: 1; }}
        .carousel-prev {{ left: 12px; }}
        .carousel-next {{ right: 12px; }}
        .carousel-btn:hover {{ background: #fff; }}
        .carousel-indicators {{ position: absolute; bottom: 12px; left: 50%; transform: translateX(-50%); display: flex; align-items: center; gap: 6px; background: rgba(0,0,0,0.5); padding: 6px 12px; border-radius: 100px; }}
        .carousel-dot {{ width: 6px; height: 6px; border-radius: 50%; background: rgba(255,255,255,0.5); cursor: pointer; transition: all 0.2s ease; }}
        .carousel-dot.active {{ background: #fff; width: 8px; height: 8px; }}
        .carousel-counter {{ font-size: 11px; color: #fff; font-weight: 600; margin-left: 6px; }}
        
        .type-section {{ margin-bottom: 64px; }}
        .type-section:last-child {{ margin-bottom: 0; }}
        .type-section-header {{ display: flex; align-items: center; gap: 12px; margin-bottom: 28px; padding-bottom: 16px; border-bottom: 1px solid var(--border); }}
        .type-section-title {{ font-size: 20px; font-weight: 600; }}
        .type-section-count {{ font-size: 13px; color: var(--text-muted); margin-left: auto; }}
        
        /* Collapsible Sections - Common */
        .collapse-icon {{ display: inline-block; width: 0; height: 0; border-left: 6px solid var(--text-muted); border-top: 5px solid transparent; border-bottom: 5px solid transparent; transition: transform 0.2s ease; margin-right: 4px; }}
        details[open] > summary .collapse-icon {{ transform: rotate(90deg); }}
        @keyframes slideDown {{ from {{ opacity: 0; transform: translateY(-8px); }} to {{ opacity: 1; transform: translateY(0); }} }}
        
        /* Collapsible Sub-sections */
        .subsection-collapsible {{ margin-bottom: 32px; }}
        .subsection-collapsible summary {{ padding: 16px 0; border-bottom: 1px solid var(--border); cursor: pointer; list-style: none; }}
        .subsection-collapsible summary::-webkit-details-marker {{ display: none; }}
        .subsection-collapsible summary::marker {{ display: none; }}
        .subsection-header {{ display: inline-flex; align-items: center; gap: 8px; width: 100%; }}
        .subsection-title {{ font-size: 18px; font-weight: 600; color: var(--text-primary); }}
        .subsection-count {{ font-size: 13px; color: var(--text-muted); margin-left: auto; padding-right: 8px; }}
        .subsection-content {{ padding-top: 24px; }}
        
        /* Collapsible Main Sections */
        .section-collapsible {{ border: none; padding: 40px 64px; border-bottom: 1px solid var(--border); }}
        .section-collapsible summary {{ cursor: pointer; list-style: none; }}
        .section-collapsible summary::-webkit-details-marker {{ display: none; }}
        .section-collapsible summary::marker {{ display: none; }}
        .section-collapsible .section-header {{ margin-bottom: 0; }}
        .section-collapsible[open] .section-header {{ margin-bottom: 48px; }}
        
        .stories-strip {{ display: flex; gap: 20px; overflow-x: auto; padding: 24px 0; }}
        .story-item {{ flex-shrink: 0; width: 140px; text-align: center; }}
        .story-thumb-wrapper {{ position: relative; width: 120px; height: 213px; margin: 0 auto 12px; border-radius: 16px; overflow: hidden; background: linear-gradient(135deg, #E1306C 0%, #F77737 50%, #FCAF45 100%); padding: 3px; }}
        .story-thumb {{ width: 100%; height: 100%; object-fit: cover; border-radius: 14px; background: var(--bg-secondary); }}
        .story-title {{ font-size: 13px; color: var(--text-primary); font-weight: 500; }}
        .story-date {{ font-size: 11px; color: var(--text-muted); margin-top: 2px; }}
        
        .interactions-grid {{ display: grid; grid-template-columns: repeat(auto-fill, minmax(380px, 1fr)); gap: 24px; }}
        .interaction-card {{ background: var(--bg-secondary); border-radius: 16px; padding: 28px; box-shadow: 0 1px 3px var(--shadow); }}
        .interaction-header {{ display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 20px; }}
        .interaction-account {{ font-weight: 600; font-size: 17px; margin-bottom: 4px; }}
        .interaction-platform {{ font-size: 13px; color: var(--text-muted); }}
        .interaction-goal-badge {{ display: inline-block; padding: 6px 14px; background: var(--bg-primary); border-radius: 100px; font-size: 12px; font-weight: 600; color: var(--text-secondary); }}
        .week-checklist {{ display: flex; gap: 8px; }}
        .day-check {{ flex: 1; text-align: center; }}
        .day-check label {{ display: block; font-size: 11px; color: var(--text-muted); margin-bottom: 8px; font-weight: 500; }}
        .day-check input[type="checkbox"] {{ appearance: none; width: 32px; height: 32px; border: 2px solid var(--border); border-radius: 8px; cursor: pointer; transition: all 0.2s ease; position: relative; }}
        .day-check input[type="checkbox"]:checked {{ background: var(--accent); border-color: var(--accent); }}
        .day-check input[type="checkbox"]:checked::after {{ content: '✓'; position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); color: #fff; font-size: 14px; font-weight: 600; }}
        
        .empty-state {{ text-align: center; padding: 48px; color: var(--text-muted); background: var(--bg-primary); border-radius: 16px; }}
        
        /* Table Search and Filter */
        .table-controls {{ display: flex; gap: 16px; margin-bottom: 16px; align-items: center; flex-wrap: wrap; }}
        .search-box {{ position: relative; flex: 1; min-width: 200px; max-width: 300px; }}
        .search-box input {{ width: 100%; padding: 10px 16px 10px 40px; border: 1px solid var(--border); border-radius: 8px; font-family: inherit; font-size: 14px; background: var(--bg-secondary); transition: border-color 0.2s ease; }}
        .search-box input:focus {{ outline: none; border-color: var(--accent); }}
        .search-box::before {{ content: '⌕'; position: absolute; left: 14px; top: 50%; transform: translateY(-50%); color: var(--text-muted); font-size: 16px; }}
        .sort-info {{ font-size: 12px; color: var(--text-muted); }}
        th.sortable {{ cursor: pointer; user-select: none; position: relative; }}
        th.sortable:hover {{ background: var(--border-light); }}
        th.sortable::after {{ content: '⇅'; margin-left: 6px; opacity: 0.4; font-size: 10px; }}
        th.sortable.asc::after {{ content: '↑'; opacity: 1; }}
        th.sortable.desc::after {{ content: '↓'; opacity: 1; }}
        .no-results {{ text-align: center; padding: 24px; color: var(--text-muted); font-size: 14px; }}
        
        /* Tab Navigation */
        .tab-nav {{ position: sticky; top: 0; z-index: 50; background: var(--bg-secondary); border-bottom: 1px solid var(--border); padding: 0 64px; }}
        .tab-list {{ display: flex; gap: 8px; overflow-x: auto; -webkit-overflow-scrolling: touch; scrollbar-width: none; }}
        .tab-list::-webkit-scrollbar {{ display: none; }}
        .tab-btn {{ padding: 16px 24px; background: none; border: none; font-family: inherit; font-size: 14px; font-weight: 500; color: var(--text-secondary); cursor: pointer; white-space: nowrap; border-bottom: 2px solid transparent; transition: all 0.2s ease; }}
        .tab-btn:hover {{ color: var(--text-primary); }}
        .tab-btn.active {{ color: var(--text-primary); border-bottom-color: var(--accent); }}
        .tab-content {{ display: none; }}
        .tab-content.active {{ display: block; }}
        
        /* Calendar Styles */
        .calendar-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 32px; }}
        .calendar-title {{ font-family: 'Montserrat', sans-serif; font-size: 28px; font-weight: 900; text-transform: uppercase; }}
        .calendar-nav {{ display: flex; gap: 8px; }}
        .calendar-nav-btn {{ width: 40px; height: 40px; display: flex; align-items: center; justify-content: center; background: var(--bg-secondary); border: 1px solid var(--border); border-radius: 8px; cursor: pointer; transition: all 0.2s ease; }}
        .calendar-nav-btn:hover {{ background: var(--bg-primary); }}
        
        /* Calendar Legend */
        .calendar-legend {{ display: flex; flex-wrap: wrap; gap: 16px; margin-bottom: 24px; padding: 16px 20px; background: var(--bg-secondary); border-radius: 12px; border: 1px solid var(--border); }}
        .calendar-legend-item {{ display: flex; align-items: center; gap: 8px; font-size: 13px; color: var(--text-secondary); }}
        .calendar-legend-icon {{ width: 20px; height: 20px; display: flex; align-items: center; justify-content: center; border-radius: 4px; font-size: 11px; }}
        .calendar-legend-icon.type-post {{ background: #F3F4F6; color: #374151; }}
        .calendar-legend-icon.type-reel {{ background: #FDF2F8; color: #BE185D; }}
        .calendar-legend-icon.type-story {{ background: #EDE9FE; color: #7C3AED; }}
        .calendar-legend-icon.type-highlight {{ background: #FEF3C7; color: #B45309; }}
        
        .view-toggle {{ display: flex; gap: 4px; background: var(--bg-primary); padding: 4px; border-radius: 8px; margin-bottom: 32px; }}
        .view-toggle-btn {{ padding: 8px 16px; background: none; border: none; font-family: inherit; font-size: 13px; font-weight: 500; color: var(--text-secondary); cursor: pointer; border-radius: 6px; transition: all 0.2s ease; }}
        .view-toggle-btn.active {{ background: var(--bg-secondary); color: var(--text-primary); box-shadow: 0 1px 3px var(--shadow); }}
        
        /* Monthly Calendar Grid */
        .calendar-grid {{ display: grid; grid-template-columns: repeat(7, 1fr); gap: 1px; background: var(--border); border-radius: 16px; overflow: hidden; }}
        .calendar-day-header {{ background: var(--bg-primary); padding: 12px 8px; text-align: center; font-size: 11px; text-transform: uppercase; letter-spacing: 1px; color: var(--text-muted); font-weight: 600; }}
        .calendar-day {{ background: var(--bg-secondary); min-height: 100px; padding: 12px; position: relative; cursor: pointer; transition: background 0.2s ease; }}
        .calendar-day:hover {{ background: var(--bg-primary); }}
        .calendar-day.other-month {{ background: var(--bg-primary); opacity: 0.5; }}
        .calendar-day.today {{ background: #FFFBEB; }}
        .calendar-day-number {{ font-size: 14px; font-weight: 600; color: var(--text-primary); margin-bottom: 8px; }}
        .calendar-day.today .calendar-day-number {{ color: var(--accent-warm); }}
        .calendar-day-posts {{ display: flex; flex-direction: column; gap: 4px; }}
        .calendar-post-indicator {{ padding: 4px 8px; border-radius: 4px; font-size: 11px; font-weight: 500; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }}
        .calendar-post-indicator.type-post {{ background: #F3F4F6; color: #374151; }}
        .calendar-post-indicator.type-reel {{ background: #FDF2F8; color: #BE185D; }}
        .calendar-post-indicator.type-highlight {{ background: #FEF3C7; color: #B45309; }}
        .calendar-post-indicator.type-story {{ background: #EDE9FE; color: #7C3AED; }}
        
        /* Weekly View - Compact Grid */
        .week-row-section {{ margin-bottom: 24px; background: var(--bg-secondary); border-radius: 16px; overflow: hidden; box-shadow: 0 1px 3px var(--shadow); }}
        .week-row-header {{ display: flex; align-items: center; gap: 16px; padding: 16px 20px; background: var(--bg-primary); border-bottom: 1px solid var(--border); }}
        .week-label {{ font-size: 12px; font-weight: 700; text-transform: uppercase; letter-spacing: 1px; color: var(--accent); }}
        .week-range {{ font-size: 15px; font-weight: 600; color: var(--text-primary); }}
        .week-count {{ font-size: 13px; color: var(--text-muted); margin-left: auto; }}
        .week-row-grid {{ display: grid; grid-template-columns: repeat(7, 1fr); gap: 1px; background: var(--border-light); }}
        .week-day-cell {{ background: var(--bg-secondary); min-height: 120px; padding: 12px; display: flex; flex-direction: column; }}
        .week-day-cell.empty {{ background: var(--bg-primary); opacity: 0.5; }}
        .week-day-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px; padding-bottom: 8px; border-bottom: 1px solid var(--border-light); }}
        .week-day-name {{ font-size: 10px; font-weight: 600; text-transform: uppercase; letter-spacing: 1px; color: var(--text-muted); }}
        .week-day-num {{ font-size: 16px; font-weight: 700; color: var(--text-primary); }}
        .week-day-items {{ display: flex; flex-direction: column; gap: 4px; flex: 1; }}
        .week-item {{ padding: 6px 8px; border-radius: 4px; font-size: 11px; font-weight: 500; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }}
        .week-item.type-post {{ background: #F3F4F6; color: #374151; }}
        .week-item.type-reel {{ background: #FDF2F8; color: #BE185D; }}
        .week-item.type-highlight {{ background: #FEF3C7; color: #B45309; }}
        .week-item.type-story {{ background: #EDE9FE; color: #7C3AED; }}
        .week-item-time {{ font-size: 10px; color: var(--text-muted); margin-right: 4px; }}
        .week-item-more {{ font-size: 10px; color: var(--text-muted); padding: 4px; }}
        
        /* Detailed Weekly View */
        .week-detail-section {{ margin-bottom: 32px; background: var(--bg-secondary); border-radius: 16px; overflow: hidden; box-shadow: 0 1px 3px var(--shadow), 0 4px 20px var(--shadow); }}
        .week-detail-header {{ display: flex; justify-content: space-between; align-items: center; padding: 20px 24px; background: linear-gradient(135deg, var(--accent) 0%, #3D3D3D 100%); color: #fff; }}
        .week-detail-title-row {{ display: flex; align-items: baseline; gap: 12px; }}
        .week-detail-label {{ font-size: 11px; font-weight: 700; text-transform: uppercase; letter-spacing: 1.5px; opacity: 0.7; }}
        .week-detail-range {{ font-size: 20px; font-weight: 600; }}
        .week-detail-stats {{ display: flex; gap: 20px; }}
        .week-stat {{ font-size: 13px; opacity: 0.8; }}
        .week-stat strong {{ font-weight: 700; }}
        .week-stat-total {{ font-size: 13px; padding: 4px 12px; background: rgba(255,255,255,0.15); border-radius: 100px; }}
        .week-detail-table {{ }}
        .week-detail-header-row {{ display: grid; grid-template-columns: 80px 70px 1fr 90px 100px 1fr; gap: 12px; padding: 14px 24px; background: var(--bg-primary); border-bottom: 1px solid var(--border); font-size: 10px; text-transform: uppercase; letter-spacing: 1.2px; color: var(--text-muted); font-weight: 600; }}
        .week-detail-row {{ display: grid; grid-template-columns: 80px 70px 1fr 90px 100px 1fr; gap: 12px; padding: 14px 24px; border-bottom: 1px solid var(--border-light); align-items: center; transition: background 0.2s ease; }}
        .week-detail-row:hover {{ background: var(--bg-primary); }}
        .week-detail-row:last-child {{ border-bottom: none; }}
        .week-detail-day {{ font-size: 13px; font-weight: 600; color: var(--text-primary); }}
        .week-detail-time {{ font-size: 13px; color: var(--text-muted); font-weight: 500; }}
        .week-detail-title {{ font-size: 14px; font-weight: 600; color: var(--text-primary); }}
        .week-detail-type {{ }}
        .week-detail-status {{ }}
        .week-detail-notes {{ font-size: 12px; color: var(--text-muted); white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }}
        
        @media (max-width: 1024px) {{
            .week-row-grid {{ grid-template-columns: repeat(4, 1fr); }}
            .week-detail-header {{ flex-direction: column; gap: 12px; align-items: flex-start; }}
            .week-detail-header-row, .week-detail-row {{ grid-template-columns: 70px 60px 1fr 80px; }}
            .week-detail-status, .week-detail-notes {{ display: none; }}
        }}
        @media (max-width: 640px) {{
            .week-row-grid {{ grid-template-columns: repeat(2, 1fr); }}
            .week-day-cell {{ min-height: 100px; }}
            .week-detail-header-row, .week-detail-row {{ grid-template-columns: 60px 1fr 70px; }}
            .week-detail-time {{ display: none; }}
        }}
        
        /* Weekly Posts Carousel - Draft Posts Section */
        .week-posts-section {{ margin-bottom: 32px; }}
        .week-posts-header {{ display: flex; align-items: center; gap: 16px; padding: 16px 20px; background: linear-gradient(135deg, var(--accent) 0%, #3D3D3D 100%); border-radius: 16px 16px 0 0; color: #fff; }}
        .week-posts-label {{ font-size: 11px; font-weight: 700; text-transform: uppercase; letter-spacing: 1.5px; opacity: 0.7; }}
        .week-posts-range {{ font-size: 18px; font-weight: 600; }}
        .week-posts-count {{ font-size: 13px; margin-left: auto; opacity: 0.8; }}
        .week-posts-carousel-wrapper {{ position: relative; background: var(--bg-secondary); border-radius: 0 0 16px 16px; box-shadow: 0 1px 3px var(--shadow), 0 4px 20px var(--shadow); }}
        .week-posts-carousel {{ display: flex; gap: 20px; overflow-x: auto; padding: 24px; scroll-snap-type: x mandatory; -webkit-overflow-scrolling: touch; scrollbar-width: thin; scrollbar-color: var(--border) var(--bg-primary); }}
        .week-posts-carousel::-webkit-scrollbar {{ height: 8px; }}
        .week-posts-carousel::-webkit-scrollbar-track {{ background: var(--bg-primary); border-radius: 4px; }}
        .week-posts-carousel::-webkit-scrollbar-thumb {{ background: var(--border); border-radius: 4px; }}
        .week-posts-carousel::-webkit-scrollbar-thumb:hover {{ background: var(--text-muted); }}
        .week-posts-carousel .post-card {{ flex-shrink: 0; width: 300px; scroll-snap-align: start; }}
        .week-posts-nav {{ position: absolute; top: 50%; transform: translateY(-50%); width: 40px; height: 40px; border-radius: 50%; background: var(--bg-secondary); border: 1px solid var(--border); display: flex; align-items: center; justify-content: center; cursor: pointer; z-index: 10; box-shadow: 0 2px 8px var(--shadow-md); transition: all 0.2s ease; font-size: 18px; color: var(--text-secondary); }}
        .week-posts-nav:hover {{ background: var(--bg-primary); color: var(--text-primary); }}
        .week-posts-nav.prev {{ left: -20px; }}
        .week-posts-nav.next {{ right: -20px; }}
        .week-posts-indicators {{ display: flex; justify-content: center; gap: 6px; padding: 16px; border-top: 1px solid var(--border-light); }}
        .week-posts-dot {{ width: 8px; height: 8px; border-radius: 50%; background: var(--border); cursor: pointer; transition: all 0.2s ease; }}
        .week-posts-dot.active {{ background: var(--accent); width: 24px; border-radius: 4px; }}
        
        @media (max-width: 1024px) {{
            .week-posts-carousel .post-card {{ width: 280px; }}
            .week-posts-nav {{ display: none; }}
        }}
        @media (max-width: 640px) {{
            .week-posts-carousel {{ gap: 16px; padding: 16px; }}
            .week-posts-carousel .post-card {{ width: 260px; }}
            .week-posts-header {{ padding: 12px 16px; border-radius: 12px 12px 0 0; }}
            .week-posts-carousel-wrapper {{ border-radius: 0 0 12px 12px; }}
        }}
        
        @media print {{ .nav {{ display: none; }} .main {{ margin-left: 0; }} .title-page {{ min-height: auto; padding: 60px; }} .tab-nav {{ display: none; }} }}
        @media (max-width: 1024px) {{ 
            .hamburger {{ display: flex; }}
            .nav {{ transform: translateX(-100%); }}
            .nav.active {{ transform: translateX(0); display: block; }}
            .main {{ margin-left: 0; }}
            .content-section {{ padding: 48px 24px; }} 
            .title-header, .title-content, .title-footer {{ padding-left: 24px; padding-right: 24px; }} 
            .title-decoration {{ display: none; }} 
            .tab-nav {{ padding: 0 16px; }}
            .calendar-day {{ min-height: 80px; padding: 8px; }}
            .calendar-post-indicator {{ font-size: 10px; padding: 2px 4px; }}
            .day-row {{ flex-direction: column; gap: 12px; }}
            .day-date {{ width: auto; display: flex; gap: 8px; align-items: baseline; }}
        }}
        @media (max-width: 640px) {{
            .calendar-day {{ min-height: 60px; }}
            .calendar-day-number {{ font-size: 12px; }}
            .calendar-post-indicator {{ display: none; }}
            .calendar-day-posts {{ display: flex; flex-direction: row; gap: 2px; }}
            .calendar-day-posts::after {{ content: attr(data-count); font-size: 10px; color: var(--text-muted); }}
            
            /* Mobile image fixes */
            .posts-grid {{ gap: 20px; }}
            .post-card-media {{ max-width: 100%; }}
            .post-card-body {{ padding: 16px; }}
            .post-card-title {{ font-size: 15px; }}
            .carousel {{ max-width: 100%; }}
            .carousel-slide img {{ max-width: 100%; height: auto; min-height: 100%; }}
            .carousel-btn {{ opacity: 1; width: 32px; height: 32px; font-size: 16px; }}
            
            /* Story thumbnails responsive scaling */
            .stories-strip {{ gap: 12px; padding: 16px 0; }}
            .story-item {{ width: 100px; }}
            .story-thumb-wrapper {{ width: 80px; height: 142px; border-radius: 12px; padding: 2px; }}
            .story-thumb {{ border-radius: 10px; max-width: 100%; }}
            .story-title {{ font-size: 11px; }}
            .story-date {{ font-size: 10px; }}
            
            /* Section padding mobile */
            .section-collapsible {{ padding: 24px 16px; }}
            .section-collapsible[open] .section-header {{ margin-bottom: 24px; }}
            .section-title {{ font-size: 28px; }}
        }}
    </style>
</head>
<body>
    <button class="hamburger" onclick="toggleNav()"><span class="bar"></span><span class="bar"></span><span class="bar"></span></button>
    <div class="nav-overlay" onclick="toggleNav()"></div>
    
    <nav class="nav">
        <div class="nav-logo"><img src="Smoothie-Bar--Logo-Assets---3D-Curved-Logo--Branding-Assets---11.21.25.png" alt="Smoothie Bar Logo"></div>
        <div class="nav-subtitle">Social Media Report</div>
        <div class="nav-section">Overview</div>
        <a href="#title">Cover</a>
        <a href="#calendar">Calendar</a>
        <a href="#schedule">Posting Schedule</a>
        <div class="nav-section">Content</div>
        <a href="#posts">Draft Posts</a>
        <a href="#posts-posts" class="sub-link">Posts</a>
        <a href="#posts-reels" class="sub-link">Reels</a>
        <a href="#posts-highlights" class="sub-link">Highlights</a>
        <div class="nav-section">Engagement</div>
        <a href="#stories">Stories</a>
        <a href="#interactions">Interactions</a>
    </nav>

    <main class="main">
        <section id="title" class="title-page">
            <img class="title-brand-image" src="Smoothie-Bar---Logo-Assets--Branding-Assets---11.21.25.png" alt="Smoothie Bar">
            <div class="title-decoration"></div>
            <header class="title-header">
                <div class="title-date">{datetime.now().strftime("%A, %B %d, %Y")}</div>
            </header>
            <div class="title-content">
                <div class="title-eyebrow">Social Media Strategy</div>
                <h1 class="title-main">Content Schedule & Asset Review</h1>
                <p class="title-description">A comprehensive overview of planned content, posting schedule, and engagement strategy.</p>
            </div>
            <footer class="title-footer">
                <div class="title-meta">
                    <div class="title-meta-item">
                        <div class="title-meta-label">Account</div>
                        <div class="title-meta-value">{config.get("AccountName", "Not Set")}</div>
                    </div>
                    <div class="title-meta-item">
                        <div class="title-meta-label">Created</div>
                        <div class="title-meta-value">{config.get("DateCreated", "Not Set")}</div>
                    </div>
                    <div class="title-meta-item">
                        <div class="title-meta-label">Total Posts</div>
                        <div class="title-meta-value">{len(posts)}</div>
                    </div>
                </div>
            </footer>
        </section>

        <details id="calendar" class="section-collapsible">
            <summary>
                <div class="section-header">
                    <div class="section-number">01</div>
                    <h2 class="section-title"><span class="collapse-icon"></span> Content Calendar</h2>
                    <p class="section-desc">Visual overview of your posting schedule.</p>
                </div>
            </summary>
            <div class="section-content">
            <div class="view-toggle">
                <button class="view-toggle-btn active" onclick="showCalendarView('monthly')">Monthly View</button>
                <button class="view-toggle-btn" onclick="showCalendarView('weekly')">Weekly View</button>
            </div>
            
            <div id="calendar-monthly" class="calendar-view">
                {monthly_calendar_html}
            </div>
            
            <div id="calendar-weekly" class="calendar-view" style="display: none;">
                {weekly_view_html}
            </div>
            
            <script>
                function showCalendarView(view) {{
                    document.querySelectorAll('.calendar-view').forEach(el => el.style.display = 'none');
                    document.querySelectorAll('.view-toggle-btn').forEach(btn => btn.classList.remove('active'));
                    document.getElementById('calendar-' + view).style.display = 'block';
                    event.target.classList.add('active');
                }}
            </script>
            </div>
        </details>

        <details id="schedule" class="section-collapsible">
            <summary>
                <div class="section-header">
                    <div class="section-number">02</div>
                    <h2 class="section-title"><span class="collapse-icon"></span> Posting Schedule</h2>
                    <p class="section-desc">Complete overview of all planned content.</p>
                </div>
            </summary>
            <div class="section-content">
            <div class="table-controls">
                <div class="search-box">
                    <input type="text" id="posts-search" placeholder="Search posts..." onkeyup="filterTable('posts-table', this.value)">
                </div>
                <span class="sort-info">Click column headers to sort</span>
            </div>
            <div class="table-wrapper">
                <table id="posts-table">
                    <thead>
                        <tr>
                            <th class="sortable" onclick="sortTable('posts-table', 0)">Title</th>
                            <th class="sortable" onclick="sortTable('posts-table', 1)">Date</th>
                            <th class="sortable" onclick="sortTable('posts-table', 2)">Time</th>
                            <th class="sortable" onclick="sortTable('posts-table', 3)">Type</th>
                            <th class="sortable" onclick="sortTable('posts-table', 4)">Status</th>
                            <th>Hashtags</th>
                            <th>Notes</th>
                        </tr>
                    </thead>
                    <tbody>
                        {"".join(f'''<tr>
                            <td><strong>{p.get("Title", "")}</strong></td>
                            <td data-sort="{convert_date_to_sortable(p.get("PostDate", ""))}">{p.get("PostDate", "")}</td>
                            <td data-sort="{convert_time_to_24hr(p.get("Time", ""))}">{p.get("Time", "")}</td>
                            <td data-sort="{p.get("Type", "post")}"><span class="type-badge type-{p.get("Type", "post").lower()}">{p.get("Type", "post")}</span></td>
                            <td data-sort="{p.get("Status", "Draft")}"><span class="status status-{p.get("Status", "Draft").lower().replace(" ", "")}">{p.get("Status", "Draft")}</span></td>
                            <td style="max-width: 200px; font-size: 13px; color: var(--text-muted);">{p.get("Hashtags", "")}</td>
                            <td style="max-width: 180px; font-size: 13px;">{p.get("Notes", "")}</td>
                        </tr>''' for p in posts) if posts else '<tr><td colspan="7" class="empty-state">No posts scheduled</td></tr>'}
                    </tbody>
                </table>
            </div>
            </div>
        </details>

        <details id="posts" class="section-collapsible">
            <summary>
                <div class="section-header">
                    <div class="section-number">03</div>
                    <h2 class="section-title"><span class="collapse-icon"></span> Draft Posts</h2>
                    <p class="section-desc">Content previews organized by format.</p>
                </div>
            </summary>
            <div class="section-content">
            
            <div class="view-toggle">
                <button class="view-toggle-btn active" onclick="showDraftsView('type')">By Type</button>
                <button class="view-toggle-btn" onclick="showDraftsView('week')">By Week</button>
            </div>
            
            <div id="drafts-by-type" class="drafts-view">
            <details id="posts-posts" class="subsection-collapsible">
                <summary>
                    <div class="subsection-header">
                        <span class="collapse-icon"></span>
                        <span class="subsection-title">Feed Posts</span>
                        <span class="subsection-count">{posts_count} posts</span>
                    </div>
                </summary>
                <div class="subsection-content">
                    <div class="posts-grid">
                        {"".join(render_post_card(p, i) for i, p in enumerate(posts) if p.get("Type", "").lower() == "post") or '<div class="empty-state">No posts scheduled</div>'}
                    </div>
                </div>
            </details>

            <details id="posts-reels" class="subsection-collapsible">
                <summary>
                    <div class="subsection-header">
                        <span class="collapse-icon"></span>
                        <span class="subsection-title">Reels</span>
                        <span class="subsection-count">{reels_count} reels</span>
                    </div>
                </summary>
                <div class="subsection-content">
                    <div class="posts-grid">
                        {"".join(render_post_card(p, i) for i, p in enumerate(posts) if p.get("Type", "").lower() == "reel") or '<div class="empty-state">No reels scheduled</div>'}
                    </div>
                </div>
            </details>

            <details id="posts-highlights" class="subsection-collapsible" {"" if highlights_count > 0 else ""}>
                <summary>
                    <div class="subsection-header">
                        <span class="collapse-icon"></span>
                        <span class="subsection-title">Highlights</span>
                        <span class="subsection-count">{highlights_count} highlights</span>
                    </div>
                </summary>
                <div class="subsection-content">
                    <div class="posts-grid">
                        {"".join(render_post_card(p, i) for i, p in enumerate(posts) if p.get("Type", "").lower() == "highlight") or '<div class="empty-state">No highlights scheduled</div>'}
                    </div>
                </div>
            </details>
            </div>
            
            <div id="drafts-by-week" class="drafts-view" style="display: none;">
                {render_posts_by_week_carousel(posts, month, year)}
            </div>
            
            <script>
                function showDraftsView(view) {{
                    document.querySelectorAll('.drafts-view').forEach(el => el.style.display = 'none');
                    document.querySelectorAll('#posts .view-toggle-btn').forEach(btn => btn.classList.remove('active'));
                    document.getElementById('drafts-by-' + view).style.display = 'block';
                    event.target.classList.add('active');
                }}
                
                function scrollWeekCarousel(btn, direction) {{
                    const wrapper = btn.closest('.week-posts-carousel-wrapper');
                    const carousel = wrapper.querySelector('.week-posts-carousel');
                    const cardWidth = carousel.querySelector('.post-card').offsetWidth + 20; // card width + gap
                    carousel.scrollBy({{ left: cardWidth * direction, behavior: 'smooth' }});
                }}
            </script>
            </div>
        </details>

        <details id="stories" class="section-collapsible">
            <summary>
                <div class="section-header">
                    <div class="section-number">04</div>
                    <h2 class="section-title"><span class="collapse-icon"></span> Stories Schedule</h2>
                    <p class="section-desc">Ephemeral content planned for the period.</p>
                </div>
            </summary>
            <div class="section-content">
            <div class="table-controls">
                <div class="search-box">
                    <input type="text" id="stories-search" placeholder="Search stories..." onkeyup="filterTable('stories-table', this.value)">
                </div>
                <span class="sort-info">Click column headers to sort</span>
            </div>
            <div class="table-wrapper">
                <table id="stories-table">
                    <thead>
                        <tr>
                            <th class="sortable" onclick="sortTable('stories-table', 0)">Title</th>
                            <th class="sortable" onclick="sortTable('stories-table', 1)">Date</th>
                            <th class="sortable" onclick="sortTable('stories-table', 2)">Time</th>
                            <th>Interactive Elements</th>
                            <th>Notes</th>
                        </tr>
                    </thead>
                    <tbody>
                        {"".join(f'''<tr>
                            <td><strong>{s.get("Title", "")}</strong></td>
                            <td data-sort="{convert_date_to_sortable(s.get("PostDate", ""))}">{s.get("PostDate", "")}</td>
                            <td data-sort="{convert_time_to_24hr(s.get("Time", ""))}">{s.get("Time", "")}</td>
                            <td>{s.get("InteractiveElements", "")}</td>
                            <td>{s.get("Notes", "")}</td>
                        </tr>''' for s in stories) if stories else '<tr><td colspan="5" class="empty-state">No stories scheduled</td></tr>'}
                    </tbody>
                </table>
            </div>
            
            <div class="type-section-header" style="margin-top: 48px;">
                <div class="type-section-title">Story Assets</div>
            </div>
            <div class="stories-strip">
                {"".join(f'''<div class="story-item">
                    <div class="story-thumb-wrapper">
                        <img class="story-thumb" src="{get_direct_image_url(s.get("MediaURL", ""))}" alt="{s.get("Title", "")}" onerror="this.style.background='#f5f5f5'">
                    </div>
                    <div class="story-title">{s.get("Title", "")}</div>
                    <div class="story-date">{s.get("PostDate", "")}</div>
                </div>''' for s in stories if s.get("MediaURL")) or '<div class="empty-state" style="width: 100%;">No story assets uploaded</div>'}
            </div>
            </div>
        </details>

        <details id="interactions" class="section-collapsible">
            <summary>
                <div class="section-header">
                    <div class="section-number">05</div>
                    <h2 class="section-title"><span class="collapse-icon"></span> Account Interactions</h2>
                    <p class="section-desc">Daily engagement targets and tracking ({len(interactions)} accounts).</p>
                </div>
            </summary>
            <div class="section-content">
            <div class="table-wrapper">
                <table>
                    <thead>
                        <tr>
                            <th>Account</th>
                            <th>Platform</th>
                            <th>Type</th>
                            <th>Goal</th>
                            <th>Mon</th>
                            <th>Tue</th>
                            <th>Wed</th>
                            <th>Thu</th>
                            <th>Fri</th>
                            <th>Sat</th>
                            <th>Sun</th>
                        </tr>
                    </thead>
                    <tbody>
                        {"".join(render_interaction_row(i) for i in interactions) if interactions else '<tr><td colspan="11" class="empty-state">No interaction targets configured</td></tr>'}
                    </tbody>
                </table>
            </div>
            </div>
        </details>
    </main>
    
    <script>
        /* Hamburger Menu */
        function toggleNav() {{
            document.querySelector('.hamburger').classList.toggle('active');
            document.querySelector('.nav').classList.toggle('active');
            document.querySelector('.nav-overlay').classList.toggle('active');
        }}
        
        // Table filtering
        function filterTable(tableId, searchText) {{
            const table = document.getElementById(tableId);
            const tbody = table.querySelector('tbody');
            const rows = tbody.querySelectorAll('tr');
            const search = searchText.toLowerCase().trim();
            let visibleCount = 0;
            
            rows.forEach(row => {{
                if (row.classList.contains('no-results-row')) {{
                    row.remove();
                    return;
                }}
                const text = row.textContent.toLowerCase();
                const match = search === '' || text.includes(search);
                row.style.display = match ? '' : 'none';
                if (match) visibleCount++;
            }});
            
            // Show "no results" message if needed
            const existingNoResults = tbody.querySelector('.no-results-row');
            if (existingNoResults) existingNoResults.remove();
            
            if (visibleCount === 0 && search !== '') {{
                const colCount = table.querySelector('thead tr').children.length;
                const noResultsRow = document.createElement('tr');
                noResultsRow.className = 'no-results-row';
                noResultsRow.innerHTML = `<td colspan="${{colCount}}" class="no-results">No results found for "${{searchText}}"</td>`;
                tbody.appendChild(noResultsRow);
            }}
        }}
        
        // Table sorting
        function sortTable(tableId, colIndex) {{
            const table = document.getElementById(tableId);
            const thead = table.querySelector('thead');
            const tbody = table.querySelector('tbody');
            const headers = thead.querySelectorAll('th');
            const rows = Array.from(tbody.querySelectorAll('tr:not(.no-results-row)'));
            
            // Determine sort direction
            const th = headers[colIndex];
            const isAsc = th.classList.contains('asc');
            
            // Clear all sort classes
            headers.forEach(h => h.classList.remove('asc', 'desc'));
            
            // Set new sort direction
            th.classList.add(isAsc ? 'desc' : 'asc');
            const direction = isAsc ? -1 : 1;
            
            // Sort rows
            rows.sort((a, b) => {{
                const aCell = a.children[colIndex];
                const bCell = b.children[colIndex];
                let aVal = aCell.dataset.sort || aCell.textContent.trim();
                let bVal = bCell.dataset.sort || bCell.textContent.trim();
                
                // Check if values look like dates (YYYY-MM-DD or HH:MM format)
                const isDateFormat = /^\\d{{4}}-\\d{{2}}-\\d{{2}}$/.test(aVal) || /^\\d{{4}}-\\d{{2}}-\\d{{2}}$/.test(bVal);
                const isTimeFormat = /^\\d{{2}}:\\d{{2}}$/.test(aVal) || /^\\d{{2}}:\\d{{2}}$/.test(bVal);
                
                // Use string comparison for dates and times (ISO format sorts correctly as strings)
                if (isDateFormat || isTimeFormat) {{
                    return aVal.localeCompare(bVal) * direction;
                }}
                
                // Try numeric comparison only if both values are purely numeric
                const aNum = parseFloat(aVal);
                const bNum = parseFloat(bVal);
                const aIsNumeric = !isNaN(aNum) && String(aNum) === aVal.trim();
                const bIsNumeric = !isNaN(bNum) && String(bNum) === bVal.trim();
                
                if (aIsNumeric && bIsNumeric) {{
                    return (aNum - bNum) * direction;
                }}
                
                // String comparison for everything else
                return aVal.localeCompare(bVal) * direction;
            }});
            
            // Re-append sorted rows
            rows.forEach(row => tbody.appendChild(row));
        }}
        
        // Carousel functions
        function getCarouselState(carousel) {{
            const container = carousel.querySelector('.carousel-container');
            const slides = carousel.querySelectorAll('.carousel-slide');
            const currentIndex = parseInt(carousel.dataset.currentIndex || 0);
            return {{ container, slides, currentIndex, total: slides.length }};
        }}
        
        function updateCarousel(carousel, newIndex) {{
            const {{ container, slides, total }} = getCarouselState(carousel);
            const index = Math.max(0, Math.min(newIndex, total - 1));
            
            carousel.dataset.currentIndex = index;
            container.style.transform = `translateX(-${{index * 100}}%)`;
            
            // Update dots
            const dots = carousel.querySelectorAll('.carousel-dot');
            dots.forEach((dot, i) => dot.classList.toggle('active', i === index));
            
            // Update counter
            const counter = carousel.querySelector('.carousel-counter');
            if (counter) counter.textContent = `${{index + 1}} / ${{total}}`;
        }}
        
        function prevSlide(btn) {{
            const carousel = btn.closest('.carousel');
            const {{ currentIndex }} = getCarouselState(carousel);
            const {{ total }} = getCarouselState(carousel);
            updateCarousel(carousel, currentIndex === 0 ? total - 1 : currentIndex - 1);
        }}
        
        function nextSlide(btn) {{
            const carousel = btn.closest('.carousel');
            const {{ currentIndex, total }} = getCarouselState(carousel);
            updateCarousel(carousel, currentIndex === total - 1 ? 0 : currentIndex + 1);
        }}
        
        function goToSlide(dot, index) {{
            const carousel = dot.closest('.carousel');
            updateCarousel(carousel, index);
        }}
        
        // Touch/swipe support for carousels
        document.addEventListener('DOMContentLoaded', () => {{
            document.querySelectorAll('.carousel').forEach(carousel => {{
                let startX = 0;
                let isDragging = false;
                
                carousel.addEventListener('touchstart', (e) => {{
                    startX = e.touches[0].clientX;
                    isDragging = true;
                }});
                
                carousel.addEventListener('touchmove', (e) => {{
                    if (!isDragging) return;
                }});
                
                carousel.addEventListener('touchend', (e) => {{
                    if (!isDragging) return;
                    isDragging = false;
                    const endX = e.changedTouches[0].clientX;
                    const diff = startX - endX;
                    
                    if (Math.abs(diff) > 50) {{
                        if (diff > 0) {{
                            nextSlide(carousel.querySelector('.carousel-next'));
                        }} else {{
                            prevSlide(carousel.querySelector('.carousel-prev'));
                        }}
                    }}
                }});
            }});
        }});
        
        // Instagram-style video controls
        function toggleVideoPlay(container) {{
            const video = container.querySelector('.video-player');
            if (!video) return;
            
            if (video.paused) {{
                video.play();
                container.classList.add('playing');
                container.classList.remove('show-pause');
            }} else {{
                video.pause();
                container.classList.remove('playing');
                container.classList.add('show-pause');
                // Remove show-pause after a moment
                setTimeout(() => container.classList.remove('show-pause'), 200);
            }}
        }}
        
        function toggleVideoMute(btn) {{
            const container = btn.closest('.video-player-container');
            const video = container.querySelector('.video-player');
            if (!video) return;
            
            video.muted = !video.muted;
            container.classList.toggle('unmuted', !video.muted);
        }}
        
        // Handle video errors - show fallback
        document.addEventListener('DOMContentLoaded', () => {{
            document.querySelectorAll('.video-player').forEach(video => {{
                video.addEventListener('error', () => {{
                    const container = video.closest('.video-player-container');
                    const fallback = container.querySelector('.video-fallback');
                    if (fallback) fallback.style.display = 'flex';
                    video.style.display = 'none';
                }});
            }});
        }});
    </script>
</body>
</html>'''
    
    return html

def render_post_card(post, index):
    """Render a single post card with carousel support for multiple images."""
    media_field = post.get("MediaURL", "")
    
    # Check for video first - includes YouTube, Vimeo, AND video file extensions
    video_platforms = ["youtube.com", "youtu.be", "vimeo.com"]
    video_extensions = [".mov", ".mp4", ".webm", ".avi", ".mkv", ".m4v"]
    
    is_platform_video = any(x in media_field.lower() for x in video_platforms)
    is_file_video = any(media_field.lower().endswith(ext) or f"{ext}?" in media_field.lower() or f"{ext}&" in media_field.lower() for ext in video_extensions)
    # Also check for Dropbox preview parameter with video files
    is_dropbox_video = "dropbox.com" in media_field.lower() and "preview=" in media_field.lower() and any(ext in media_field.lower() for ext in video_extensions)
    
    is_video = is_platform_video or is_file_video or is_dropbox_video
    
    if is_video:
        if is_platform_video:
            embed_url = get_embed_url(media_field)
            media_html = f'<div class="post-card-media video-container"><iframe src="{embed_url}" allowfullscreen></iframe></div>'
        else:
            # For direct video files (Dropbox, etc.), use HTML5 video player
            video_url = get_direct_video_url(media_field)
            filename = extract_filename_from_url(media_field)
            # Determine video type from extension
            video_type = "video/mp4"  # Default
            if ".mov" in media_field.lower():
                video_type = "video/quicktime"
            elif ".webm" in media_field.lower():
                video_type = "video/webm"
            elif ".avi" in media_field.lower():
                video_type = "video/x-msvideo"
            
            media_html = f'''<div class="post-card-media video-player-container" onclick="toggleVideoPlay(this)">
                <video class="video-player" playsinline preload="metadata" muted loop poster="">
                    <source src="{video_url}" type="{video_type}">
                    <source src="{video_url}" type="video/mp4">
                    Your browser does not support the video tag.
                </video>
                <div class="video-play-indicator">
                    <svg viewBox="0 0 24 24" fill="white"><polygon points="5 3 19 12 5 21 5 3"/></svg>
                </div>
                <button class="video-mute-btn" onclick="event.stopPropagation(); toggleVideoMute(this)">
                    <svg class="mute-icon" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polygon points="11 5 6 9 2 9 2 15 6 15 11 19 11 5"/><line x1="23" y1="9" x2="17" y2="15"/><line x1="17" y1="9" x2="23" y2="15"/></svg>
                    <svg class="unmute-icon" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polygon points="11 5 6 9 2 9 2 15 6 15 11 19 11 5"/><path d="M19.07 4.93a10 10 0 0 1 0 14.14M15.54 8.46a5 5 0 0 1 0 7.07"/></svg>
                </button>
                <div class="video-fallback" onclick="event.stopPropagation(); window.open('{media_field}', '_blank')">
                    <div class="video-placeholder-content">
                        <svg width="48" height="48" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5">
                            <polygon points="5 3 19 12 5 21 5 3" fill="currentColor"/>
                        </svg>
                        <span class="video-placeholder-text">Click to open video</span>
                        <span class="video-placeholder-filename">{filename}</span>
                    </div>
                </div>
            </div>'''
    else:
        # Parse multiple URLs (carousel support)
        media_urls = parse_media_urls(media_field)
        
        if len(media_urls) > 1:
            # Carousel with multiple images - add referrerpolicy for mobile Dropbox support
            slides_html = ''.join(f'''<div class="carousel-slide" data-index="{i}">
                <img src="{url}" alt="{post.get("Title", "")} - Image {i+1}" referrerpolicy="no-referrer" crossorigin="anonymous" onerror="this.parentElement.innerHTML='<div class=\\'no-media\\'>Image not available</div>'">
            </div>''' for i, url in enumerate(media_urls))
            
            dots_html = ''.join(f'<span class="carousel-dot{" active" if i == 0 else ""}" data-index="{i}" onclick="goToSlide(this, {i})"></span>' for i in range(len(media_urls)))
            
            media_html = f'''<div class="carousel" data-card="{index}">
                <div class="carousel-container">
                    {slides_html}
                </div>
                <button class="carousel-btn carousel-prev" onclick="prevSlide(this)">‹</button>
                <button class="carousel-btn carousel-next" onclick="nextSlide(this)">›</button>
                <div class="carousel-indicators">
                    {dots_html}
                    <span class="carousel-counter">1 / {len(media_urls)}</span>
                </div>
            </div>'''
        elif len(media_urls) == 1:
            # Single image - add referrerpolicy for better mobile Dropbox support
            media_html = f'<img class="post-card-media" src="{media_urls[0]}" alt="{post.get("Title", "")}" referrerpolicy="no-referrer" crossorigin="anonymous" onerror="this.outerHTML=\'<div class=\\\'no-media\\\'>Image not available</div>\'">'
        else:
            media_html = '<div class="no-media">No media uploaded</div>'
    
    # Format date and time for display
    post_date = post.get("PostDate", "")
    post_time = post.get("Time", "")
    date_display = f"{post_date} • {post_time}" if post_date and post_time else post_date or post_time or ""
    
    return f'''<div class="post-card" id="post-{index}">
        {media_html}
        <div class="post-card-body">
            <div class="post-card-header">
                <span class="post-card-title">{post.get("Title", "Untitled")}</span>
                <span class="type-badge type-{post.get("Type", "post").lower()}">{post.get("Type", "post")}</span>
            </div>
            {f'<div class="post-card-date">{date_display}</div>' if date_display else ''}
            <p class="post-card-caption">{post.get("Caption", "No caption provided")}</p>
            <div class="post-card-hashtags">{post.get("Hashtags", "")}</div>
        </div>
    </div>'''

def render_interaction_card(interaction):
    """Render a single interaction card."""
    days = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
    checkboxes = "".join(f'''<div class="day-check">
        <label>{day}</label>
        <input type="checkbox" {"checked" if interaction.get(day) == "TRUE" else ""}>
    </div>''' for day in days)
    
    return f'''<div class="interaction-card">
        <div class="interaction-header">
            <div>
                <div class="interaction-account">@{interaction.get("AccountName", "")}</div>
                <div class="interaction-platform">{interaction.get("Platform", "")} · {interaction.get("InteractionType", "")}</div>
            </div>
            <div class="interaction-goal-badge">{interaction.get("DailyGoal", "0")}/day</div>
        </div>
        <div class="week-checklist">{checkboxes}</div>
    </div>'''


def render_interaction_row(interaction):
    """Render a single interaction as a table row (compact view for large lists)."""
    days = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
    day_cells = "".join(f'''<td style="text-align: center;">
        <input type="checkbox" style="width: 18px; height: 18px; cursor: pointer;" {"checked" if interaction.get(day) == "TRUE" else ""}>
    </td>''' for day in days)
    
    return f'''<tr>
        <td><strong>@{interaction.get("AccountName", "")}</strong></td>
        <td style="font-size: 13px; color: var(--text-muted);">{interaction.get("Platform", "").replace("[", "").replace("]", "")}</td>
        <td style="font-size: 13px;">{interaction.get("InteractionType", "").replace("[", "").replace("]", "")}</td>
        <td style="font-weight: 600;">{interaction.get("DailyGoal", "0").replace("[", "").replace("]", "")}/day</td>
        {day_cells}
    </tr>'''

def render_monthly_calendar(posts, stories, month, year):
    """Render the monthly calendar grid HTML with posts AND stories."""
    cal = calendar.Calendar(firstweekday=6)  # Sunday start
    month_name = calendar.month_name[month]
    today = datetime.now()
    
    # Build a dict of day -> items for quick lookup (combine posts + stories)
    items_by_day = {}
    
    # Add posts
    for post in posts:
        post_date = parse_date(post.get("PostDate", ""), year)
        if post_date and post_date.month == month and post_date.year == year:
            day = post_date.day
            if day not in items_by_day:
                items_by_day[day] = []
            item = post.copy()
            item["_content_type"] = "post"
            items_by_day[day].append(item)
    
    # Add stories
    for story in stories:
        story_date = parse_date(story.get("PostDate", ""), year)
        if story_date and story_date.month == month and story_date.year == year:
            day = story_date.day
            if day not in items_by_day:
                items_by_day[day] = []
            item = story.copy()
            item["_content_type"] = "story"
            item["Type"] = "story"  # Mark as story type for styling
            items_by_day[day].append(item)
    
    # Day headers
    day_headers = ['Sun', 'Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat']
    headers_html = ''.join(f'<div class="calendar-day-header">{d}</div>' for d in day_headers)
    
    # Calendar days
    days_html = ''
    for day in cal.itermonthdays(year, month):
        if day == 0:
            days_html += '<div class="calendar-day other-month"></div>'
        else:
            is_today = (today.year == year and today.month == month and today.day == day)
            today_class = ' today' if is_today else ''
            day_items = items_by_day.get(day, [])
            
            posts_indicators = ''
            for p in day_items[:3]:  # Show max 3 indicators
                post_type = p.get("Type", "post").lower()
                title = p.get("Title", "")[:20]
                posts_indicators += f'<div class="calendar-post-indicator type-{post_type}">{title}</div>'
            
            if len(day_items) > 3:
                posts_indicators += f'<div class="calendar-post-indicator" style="background: var(--bg-primary); color: var(--text-muted);">+{len(day_items) - 3} more</div>'
            
            days_html += f'''<div class="calendar-day{today_class}" data-day="{day}">
                <div class="calendar-day-number">{day}</div>
                <div class="calendar-day-posts" data-count="{len(day_items)}">{posts_indicators}</div>
            </div>'''
    
    return f'''<div class="calendar-legend">
        <div class="calendar-legend-item">
            <div class="calendar-legend-icon type-post">
                <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><rect x="3" y="3" width="18" height="18" rx="3"/><circle cx="8.5" cy="8.5" r="1.5" fill="currentColor"/><path d="M21 15l-5-5L5 21"/></svg>
            </div>
            <span>Post</span>
        </div>
        <div class="calendar-legend-item">
            <div class="calendar-legend-icon type-reel">
                <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><polygon points="5 3 19 12 5 21 5 3" fill="currentColor"/></svg>
            </div>
            <span>Reel</span>
        </div>
        <div class="calendar-legend-item">
            <div class="calendar-legend-icon type-story">
                <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><circle cx="12" cy="12" r="9"/><circle cx="12" cy="12" r="4" fill="currentColor"/></svg>
            </div>
            <span>Story</span>
        </div>
        <div class="calendar-legend-item">
            <div class="calendar-legend-icon type-highlight">
                <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5"><polygon points="12 2 15 9 22 9 17 14 19 21 12 17 5 21 7 14 2 9 9 9" fill="currentColor"/></svg>
            </div>
            <span>Highlight</span>
        </div>
    </div>
    <div class="calendar-header">
        <h2 class="calendar-title">{month_name} {year}</h2>
    </div>
    <div class="calendar-grid">
        {headers_html}
        {days_html}
    </div>'''


def render_weekly_view(posts, stories, month, year):
    """Render a detailed weekly breakdown showing all posts and stories by week."""
    # Combine posts and stories
    all_items = []
    for post in posts:
        item = post.copy()
        item["_content_type"] = "post"
        all_items.append(item)
    for story in stories:
        item = story.copy()
        item["_content_type"] = "story"
        item["Type"] = "story"
        item["Status"] = "Draft"
        all_items.append(item)
    
    # Build items by day for quick lookup
    items_by_day = {}
    for item in all_items:
        item_date = parse_date(item.get("PostDate", ""), year)
        if item_date and item_date.month == month and item_date.year == year:
            day = item_date.day
            if day not in items_by_day:
                items_by_day[day] = []
            items_by_day[day].append(item)
    
    # Sort items within each day by time
    for day in items_by_day:
        items_by_day[day].sort(key=lambda x: x.get("Time", "99:99"))
    
    if not items_by_day:
        return '<div class="empty-state">No content scheduled for this month</div>'
    
    # Get calendar structure
    cal = calendar.Calendar(firstweekday=6)  # Sunday start
    month_days = list(cal.itermonthdays(year, month))
    day_names = ['Sun', 'Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat']
    
    # Generate week sections
    weeks_html = ''
    week_num = 0
    
    i = 0
    while i < len(month_days):
        week_days = month_days[i:i+7]
        i += 7
        
        # Skip weeks with no days in this month
        if all(d == 0 for d in week_days):
            continue
        
        week_num += 1
        
        # Calculate week date range
        valid_days = [d for d in week_days if d > 0]
        if not valid_days:
            continue
            
        start_day = min(valid_days)
        end_day = max(valid_days)
        month_abbr = calendar.month_abbr[month]
        date_range = f"{month_abbr} {start_day}-{end_day}"
        
        # Collect all items for this week
        week_items = []
        for idx, day in enumerate(week_days):
            if day > 0 and day in items_by_day:
                for item in items_by_day[day]:
                    week_items.append((day, idx, item))
        
        # Skip weeks with no content
        if not week_items:
            continue
        
        # Count by type
        post_count = sum(1 for _, _, item in week_items if item.get("Type", "").lower() in ["post", "reel", "highlight"])
        story_count = sum(1 for _, _, item in week_items if item.get("Type", "").lower() == "story")
        
        # Build day rows
        days_html = ''
        current_day = None
        
        for day, day_idx, item in sorted(week_items, key=lambda x: (x[0], x[2].get("Time", "99:99"))):
            day_date = datetime(year, month, day)
            weekday_name = day_names[day_idx]
            item_type = item.get("Type", "post").lower()
            title = item.get("Title", "Untitled")
            time = item.get("Time", "")
            status = item.get("Status", "Draft")
            notes = item.get("Notes", "") or item.get("InteractiveElements", "")
            
            # Show day label only for first item of each day
            if day != current_day:
                day_label = f'<div class="week-detail-day">{weekday_name} {day}</div>'
                current_day = day
            else:
                day_label = '<div class="week-detail-day"></div>'
            
            days_html += f'''<div class="week-detail-row">
                {day_label}
                <div class="week-detail-time">{time}</div>
                <div class="week-detail-title">{title}</div>
                <div class="week-detail-type"><span class="type-badge type-{item_type}">{item_type}</span></div>
                <div class="week-detail-status"><span class="status status-{status.lower().replace(" ", "")}">{status}</span></div>
                <div class="week-detail-notes">{notes[:40]}{"..." if len(notes) > 40 else ""}</div>
            </div>'''
        
        weeks_html += f'''<div class="week-detail-section">
            <div class="week-detail-header">
                <div class="week-detail-title-row">
                    <span class="week-detail-label">Week {week_num}</span>
                    <span class="week-detail-range">{date_range}</span>
                </div>
                <div class="week-detail-stats">
                    <span class="week-stat"><strong>{post_count}</strong> posts</span>
                    <span class="week-stat"><strong>{story_count}</strong> stories</span>
                    <span class="week-stat-total"><strong>{len(week_items)}</strong> total</span>
                </div>
            </div>
            <div class="week-detail-table">
                <div class="week-detail-header-row">
                    <div class="week-detail-day">Day</div>
                    <div class="week-detail-time">Time</div>
                    <div class="week-detail-title">Title</div>
                    <div class="week-detail-type">Type</div>
                    <div class="week-detail-status">Status</div>
                    <div class="week-detail-notes">Notes</div>
                </div>
                {days_html}
            </div>
        </div>'''
    
    return weeks_html


def render_posts_by_week_carousel(posts, month, year):
    """
    Render posts grouped by week as horizontal carousels.
    Each week gets its own swipeable carousel of post cards.
    """
    # Build posts by day for quick lookup (only posts, not stories)
    posts_by_day = {}
    for i, post in enumerate(posts):
        post_date = parse_date(post.get("PostDate", ""), year)
        if post_date and post_date.month == month and post_date.year == year:
            day = post_date.day
            if day not in posts_by_day:
                posts_by_day[day] = []
            posts_by_day[day].append((i, post))
    
    # Sort posts within each day by time
    for day in posts_by_day:
        posts_by_day[day].sort(key=lambda x: x[1].get("Time", "99:99"))
    
    if not posts_by_day:
        return '<div class="empty-state">No posts scheduled for this month</div>'
    
    # Get calendar structure
    cal = calendar.Calendar(firstweekday=6)  # Sunday start
    month_days = list(cal.itermonthdays(year, month))
    
    # Generate week sections
    weeks_html = ''
    week_num = 0
    
    i = 0
    while i < len(month_days):
        week_days = month_days[i:i+7]
        i += 7
        
        # Skip weeks with no days in this month
        if all(d == 0 for d in week_days):
            continue
        
        week_num += 1
        
        # Calculate week date range
        valid_days = [d for d in week_days if d > 0]
        if not valid_days:
            continue
            
        start_day = min(valid_days)
        end_day = max(valid_days)
        month_abbr = calendar.month_abbr[month]
        date_range = f"{month_abbr} {start_day}-{end_day}"
        
        # Collect all posts for this week
        week_posts = []
        for day in week_days:
            if day > 0 and day in posts_by_day:
                for idx, post in posts_by_day[day]:
                    week_posts.append((day, idx, post))
        
        # Skip weeks with no posts
        if not week_posts:
            continue
        
        # Sort by day and time
        week_posts.sort(key=lambda x: (x[0], x[2].get("Time", "99:99")))
        
        # Generate post cards HTML for this week
        cards_html = ''
        for day, idx, post in week_posts:
            cards_html += render_post_card(post, idx)
        
        weeks_html += f'''<div class="week-posts-section">
            <div class="week-posts-header">
                <span class="week-posts-label">Week {week_num}</span>
                <span class="week-posts-range">{date_range}</span>
                <span class="week-posts-count">{len(week_posts)} posts</span>
            </div>
            <div class="week-posts-carousel-wrapper">
                <button class="week-posts-nav prev" onclick="scrollWeekCarousel(this, -1)">‹</button>
                <div class="week-posts-carousel">
                    {cards_html}
                </div>
                <button class="week-posts-nav next" onclick="scrollWeekCarousel(this, 1)">›</button>
            </div>
        </div>'''
    
    return weeks_html


def convert_date_to_sortable(date_str, year=None):
    """Convert date string to sortable ISO format (YYYY-MM-DD) for proper sorting."""
    if not date_str:
        return "9999-99-99"  # Default to end for empty dates
    
    parsed = parse_date(date_str, year if year else datetime.now().year)
    if parsed:
        return parsed.strftime("%Y-%m-%d")
    
    return date_str  # Return as-is if couldn't parse


def convert_time_to_24hr(time_str):
    """Convert time string like '9:00 AM' or '12:30 PM' to 24-hour format for sorting."""
    if not time_str:
        return "99:99"  # Default to end for empty times
    
    time_str = time_str.strip().upper()
    
    # Try to parse AM/PM format
    try:
        # Handle formats like "9:00 AM", "12:30 PM", "9:00AM", "9AM"
        import re
        match = re.match(r'(\d{1,2}):?(\d{2})?\s*(AM|PM)?', time_str, re.IGNORECASE)
        if match:
            hour = int(match.group(1))
            minute = int(match.group(2)) if match.group(2) else 0
            period = match.group(3)
            
            if period:
                period = period.upper()
                if period == 'PM' and hour != 12:
                    hour += 12
                elif period == 'AM' and hour == 12:
                    hour = 0
            
            return f"{hour:02d}:{minute:02d}"
    except:
        pass
    
    return time_str  # Return as-is if couldn't parse


def get_direct_video_url(url):
    """
    Get a direct video URL for Dropbox and other hosts that can be used
    with HTML5 <video> element.
    
    For Dropbox, converts to dl.dropboxusercontent.com with raw=1 for direct streaming.
    """
    if not url:
        return url
    
    url = url.strip()
    from urllib.parse import unquote
    
    # Handle Dropbox URLs - convert to direct streaming URL
    if "dropbox.com" in url:
        # Special handling for folder preview URLs (with preview= parameter)
        # The filename is in the preview= parameter, not the path!
        # We need to reconstruct the URL with the filename in the path
        if "preview=" in url and "/scl/fo/" in url:
            # Extract the preview filename
            preview_match = re.search(r'preview=([^&]+)', url)
            if preview_match:
                preview_filename = unquote(preview_match.group(1).replace('+', ' '))
                
                # Extract the folder path (everything before the ?)
                base_path = url.split('?')[0]
                
                # Extract rlkey parameter (needed for access)
                rlkey_match = re.search(r'rlkey=([^&]+)', url)
                rlkey = rlkey_match.group(1) if rlkey_match else ""
                
                # Construct new URL with filename appended to path
                # Format: https://dl.dropboxusercontent.com/scl/fo/{folder}/{subfolder}/{filename}?rlkey=xxx&raw=1
                new_url = base_path.rstrip('/')
                # URL encode the filename for the path
                from urllib.parse import quote
                encoded_filename = quote(preview_filename)
                new_url = f"{new_url}/{encoded_filename}"
                
                # Convert to direct download domain
                new_url = new_url.replace("www.dropbox.com", "dl.dropboxusercontent.com")
                new_url = new_url.replace("dropbox.com", "dl.dropboxusercontent.com")
                
                # Add required parameters
                new_url = f"{new_url}?raw=1"
                if rlkey:
                    new_url = f"{new_url}&rlkey={rlkey}"
                
                return new_url
        
        # Standard file links (no preview parameter needed)
        if "www.dropbox.com" in url:
            url = url.replace("www.dropbox.com", "dl.dropboxusercontent.com")
        elif "dl.dropboxusercontent.com" not in url:
            url = url.replace("dropbox.com", "dl.dropboxusercontent.com")
        
        # Remove the session token (st=) which expires
        url = re.sub(r'[&?]st=[^&]*', '', url)
        url = url.replace('&&', '&').replace('?&', '?')
        url = url.rstrip('&').rstrip('?')
        
        # Remove e= parameter 
        url = re.sub(r'[&?]e=[^&]*', '', url)
        url = url.replace('&&', '&').replace('?&', '?')
        url = url.rstrip('&').rstrip('?')
        
        # Ensure raw=1 for direct streaming
        if "dl=0" in url:
            url = url.replace("dl=0", "raw=1")
        elif "dl=1" in url:
            url = url.replace("dl=1", "raw=1")
        elif "raw=" not in url:
            if "?" in url:
                url = url + "&raw=1"
            else:
                url = url + "?raw=1"
    
    return url


def extract_filename_from_url(url):
    """
    Extract filename from a URL, handling encoded characters.
    """
    if not url:
        return "Video"
    
    try:
        from urllib.parse import unquote
        
        # Look for filename in preview= parameter first (Dropbox)
        if "preview=" in url:
            preview_match = re.search(r'preview=([^&]+)', url)
            if preview_match:
                filename = unquote(preview_match.group(1).replace('+', ' '))
                # Truncate if too long
                if len(filename) > 40:
                    filename = filename[:37] + "..."
                return filename
        
        # Look for filename in path
        path_parts = url.split('/')
        for part in reversed(path_parts):
            if part and '.' in part:
                filename = unquote(part.split('?')[0].replace('+', ' '))
                if len(filename) > 40:
                    filename = filename[:37] + "..."
                return filename
        
        return "Video"
    except:
        return "Video"


def get_embed_url(url):
    """Convert video URLs to embed format."""
    if "youtube.com/watch" in url:
        video_id = url.split("v=")[1].split("&")[0] if "v=" in url else ""
        return f"https://www.youtube.com/embed/{video_id}"
    elif "youtu.be/" in url:
        video_id = url.split("youtu.be/")[1].split("?")[0]
        return f"https://www.youtube.com/embed/{video_id}"
    elif "vimeo.com/" in url:
        video_id = url.split("vimeo.com/")[1].split("?")[0]
        return f"https://player.vimeo.com/video/{video_id}"
    return url


def convert_imgur_url(url):
    """
    Convert Imgur page URLs to direct image URLs.
    
    Converts:
        https://imgur.com/zc7nfaj  →  https://i.imgur.com/zc7nfaj.jpg
        https://imgur.com/a/ABC123  →  (album - returns first item or placeholder)
        https://i.imgur.com/xyz.jpg  →  unchanged (already direct)
    """
    if not url:
        return url
    
    url = url.strip()
    
    # Already a direct image URL (i.imgur.com with extension)
    if "i.imgur.com/" in url and any(url.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']):
        return url
    
    # Handle album URLs - these can't be directly converted without API
    # Return a note or try first image
    if "imgur.com/a/" in url or "imgur.com/gallery/" in url:
        # Album URLs need manual conversion - extract ID and try common pattern
        # For albums, you'd need to get individual image IDs from the album
        return url  # Can't auto-convert albums, return as-is
    
    # Single image page URL: https://imgur.com/zc7nfaj
    # Convert to: https://i.imgur.com/zc7nfaj.jpg
    if "imgur.com/" in url and "i.imgur.com" not in url:
        # Extract the image ID
        parts = url.split("imgur.com/")
        if len(parts) > 1:
            image_id = parts[1].split("?")[0].split("#")[0].strip("/")
            # Skip if it looks like an album or gallery path
            if image_id and "/" not in image_id and len(image_id) >= 5:
                return f"https://i.imgur.com/{image_id}.jpg"
    
    return url


def convert_dropbox_url(url):
    """
    Convert Dropbox share URLs to direct download URLs that work on mobile.
    
    Converts:
        https://www.dropbox.com/...?dl=0  →  https://dl.dropboxusercontent.com/...?raw=1
        https://www.dropbox.com/scl/fi/...  →  https://dl.dropboxusercontent.com/scl/fi/...?raw=1
    
    Using dl.dropboxusercontent.com with raw=1 avoids redirect issues on mobile browsers.
    
    IMPORTANT: Removes the &st= session token parameter which EXPIRES and causes
    mobile loading failures. Only the rlkey= parameter is needed for permanent access.
    """
    if not url:
        return url
    
    url = url.strip()
    
    # Check if it's a Dropbox URL
    if "dropbox.com" not in url:
        return url
    
    # Convert www.dropbox.com to dl.dropboxusercontent.com for direct access
    if "www.dropbox.com" in url:
        url = url.replace("www.dropbox.com", "dl.dropboxusercontent.com")
    elif "dropbox.com" in url and "dl.dropboxusercontent.com" not in url:
        url = url.replace("dropbox.com", "dl.dropboxusercontent.com")
    
    # CRITICAL: Remove the &st= session token parameter - it EXPIRES and breaks mobile loading!
    # The rlkey= parameter is the permanent share key, st= is a temporary session token
    import re
    url = re.sub(r'[&?]st=[^&]*', '', url)
    
    # Clean up any double && or trailing & or ?& that might result from removal
    url = url.replace('&&', '&').replace('?&', '?')
    url = url.rstrip('&').rstrip('?')
    
    # Use raw=1 instead of dl=1 for better mobile compatibility
    if "dl=0" in url:
        url = url.replace("dl=0", "raw=1")
    elif "dl=1" in url:
        url = url.replace("dl=1", "raw=1")
    elif "raw=" not in url:
        if "?" in url:
            url = url + "&raw=1"
        else:
            url = url + "?raw=1"
    
    return url


def get_direct_image_url(url):
    """
    Get a direct image URL that can be used in <img src="">.
    Handles Dropbox, Imgur, Google Drive, and other common hosts.
    """
    if not url:
        return url
    
    url = url.strip()
    
    # Handle Dropbox (check first - most common now)
    if "dropbox.com" in url:
        return convert_dropbox_url(url)
    
    # Handle Imgur
    if "imgur.com" in url:
        return convert_imgur_url(url)
    
    # Handle Google Drive
    # Convert: https://drive.google.com/file/d/FILE_ID/view
    # To: https://drive.google.com/uc?export=view&id=FILE_ID
    if "drive.google.com/file/d/" in url:
        try:
            file_id = url.split("/d/")[1].split("/")[0]
            return f"https://drive.google.com/uc?export=view&id={file_id}"
        except:
            pass
    
    # Already a direct URL or unchanged
    return url


def parse_media_urls(media_field):
    """
    Parse a media URL field that may contain multiple URLs.
    Returns a list of direct image URLs.
    
    Handles:
        - Single URL
        - Multiple URLs separated by newlines
        - URLs with descriptive text (e.g., "https://... - Album")
    """
    if not media_field:
        return []
    
    urls = []
    
    # Split by newlines
    lines = media_field.strip().split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Extract URL from line (may have description like "https://... - Album")
        # Look for URL pattern
        url_match = re.search(r'(https?://[^\s]+)', line)
        if url_match:
            url = url_match.group(1)
            # Clean up any trailing description markers
            url = url.split(' - ')[0].strip()
            url = url.split(' ')[0].strip()  # Take only the URL part
            
            # Convert to direct URL
            direct_url = get_direct_image_url(url)
            
            # Skip album URLs that couldn't be converted
            if "imgur.com/a/" not in direct_url and "imgur.com/gallery/" not in direct_url:
                urls.append(direct_url)
    
    return urls

def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_report.py input.docx [output.html]")
        print("\nThis script reads a Word document and generates an HTML report.")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace(".docx", "_report.html")
    
    if not os.path.exists(input_file):
        print(f"Error: File '{input_file}' not found.")
        sys.exit(1)
    
    print(f"📄 Reading: {input_file}")
    
    # Load document
    doc = Document(input_file)
    tables = extract_tables(doc)
    tables_with_hyperlinks = extract_tables_with_hyperlinks(doc)
    
    print(f"   Found {len(tables)} tables")
    
    # Identify tables using hyperlink-aware extraction for posts schedule
    # This ensures Photo Links column gets the actual URLs from hyperlinks
    identified = identify_tables(tables_with_hyperlinks, doc)
    
    config = parse_config(tables)
    print(f"   Config: {config}")
    
    # Parse posts from schedule table (using hyperlink-aware table)
    posts = []
    if identified["posts_schedule"]:
        posts = parse_posts_table(identified["posts_schedule"])
        print(f"   Posts from schedule: {len(posts)}")
    
    # Merge with post detail blocks (for captions/media)
    # Use tables_with_hyperlinks instead of regular tables to get embedded URLs
    # Re-identify post_blocks using the hyperlink-aware tables
    identified_with_hyperlinks = identify_tables(tables_with_hyperlinks, doc)
    
    post_blocks = []
    for block in identified_with_hyperlinks["post_blocks"]:
        block_data = {}
        post_type = "post"
        for row in block:
            if len(row) >= 2:
                label = row[0].strip().lower()
                value = row[1].strip()
                if value.startswith("[") and value.endswith("]"):
                    continue
                if label == "title":
                    block_data["Title"] = value
                elif label in ["caption", "description"]:
                    block_data["Caption"] = value
                elif label == "hashtags":
                    block_data["Hashtags"] = value
                elif "video" in label:
                    block_data["MediaURL"] = value
                    post_type = "reel"
                elif "cover" in label:
                    block_data["MediaURL"] = value
                    post_type = "highlight"
                elif "image" in label or "url" in label:
                    block_data["MediaURL"] = value
        if block_data.get("Title"):
            block_data["Type"] = post_type
            post_blocks.append(block_data)
    
    # Merge post blocks with schedule
    for block in post_blocks:
        title = block.get("Title", "")
        matched = False
        for post in posts:
            if post.get("Title") == title:
                post.update({k: v for k, v in block.items() if v})
                matched = True
                break
        if not matched:
            posts.append(block)
    
    print(f"   Total posts after merge: {len(posts)}")
    
    # Parse stories
    stories = []
    if identified["stories"]:
        stories = parse_stories_table(identified["stories"])
    print(f"   Stories: {len(stories)}")
    
    # Parse interactions
    interactions = []
    if identified["interactions"]:
        interactions = parse_interactions_table(identified["interactions"])
    print(f"   Interactions: {len(interactions)}")
    
    # Generate HTML
    html = generate_html(config, posts, stories, interactions)
    
    # Write output
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(html)
    
    print(f"\n✅ Report generated: {output_file}")
    print(f"   Open in any browser to view!")

if __name__ == "__main__":
    main()
