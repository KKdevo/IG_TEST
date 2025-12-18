import sys
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# Use command line arg or default
doc_path = sys.argv[1] if len(sys.argv) > 1 else 'Content_Schedule_INPUT/Content-Schedule-Template-DEC.docx'
doc = Document(doc_path)
t = doc.tables[1]

# Get the relationship part for hyperlinks
rels = doc.part.rels

print(f"Checking Photo Links column (index 3) in: {doc_path}\n")
for i, row in enumerate(t.rows[1:]):  # Check ALL rows
    cell = row.cells[3]
    title = row.cells[0].text[:35].replace('\n', ' ')
    display = cell.text.strip().replace('\n', ' ')[:30]
    
    print(f"Row {i+1}: Title='{title}' | Display='{display}'")
    
    # Check for hyperlinks in cell
    found_links = []
    for p in cell.paragraphs:
        for hl in p._element.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink'):
            r_id = hl.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if r_id and r_id in rels:
                found_links.append(rels[r_id].target_ref)
    
    if found_links:
        for link in found_links:
            print(f"   -> {link[:80]}...")
    else:
        print(f"   -> NO HYPERLINKS FOUND")
    print()
