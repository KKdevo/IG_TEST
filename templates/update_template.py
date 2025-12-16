#!/usr/bin/env python3
"""
Update Template Script
Adds 'Photo Links' and 'Caption' columns to the Content Schedule template.

Usage:
    python update_template.py
"""

from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy
import re

def add_columns_to_table(table, insert_after_col, new_col_headers):
    """
    Insert new columns into a table after a specific column index.
    Uses direct XML manipulation for reliable column insertion.
    """
    tbl = table._tbl
    
    # 1. Update tblGrid to add column definitions
    tblGrid = tbl.find(qn('w:tblGrid'))
    gridCols = tblGrid.findall(qn('w:gridCol'))
    
    # Copy width from reference column
    ref_col = gridCols[insert_after_col] if insert_after_col < len(gridCols) else gridCols[-1]
    ref_width = ref_col.get(qn('w:w'))
    
    # Insert new gridCol elements
    for i, header in enumerate(new_col_headers):
        new_gridCol = OxmlElement('w:gridCol')
        if ref_width:
            new_gridCol.set(qn('w:w'), ref_width)
        # Insert after the reference position
        insert_pos = list(tblGrid).index(gridCols[insert_after_col]) + 1 + i
        tblGrid.insert(insert_pos, new_gridCol)
    
    # 2. Add cells to each row
    for row_idx, row in enumerate(table.rows):
        tr = row._tr
        tcs = tr.findall(qn('w:tc'))
        
        if insert_after_col < len(tcs):
            ref_tc = tcs[insert_after_col]
            insert_idx = list(tr).index(ref_tc) + 1
        else:
            ref_tc = tcs[-1]
            insert_idx = len(list(tr))
        
        # Create new cells
        for i, header in enumerate(new_col_headers):
            # Create new tc element
            new_tc = OxmlElement('w:tc')
            
            # Copy cell properties from reference
            ref_tcPr = ref_tc.find(qn('w:tcPr'))
            if ref_tcPr is not None:
                new_tcPr = deepcopy(ref_tcPr)
                new_tc.append(new_tcPr)
            
            # Add paragraph with text
            new_p = OxmlElement('w:p')
            
            if row_idx == 0:
                # Header row - add text
                new_r = OxmlElement('w:r')
                new_t = OxmlElement('w:t')
                new_t.text = header
                new_r.append(new_t)
                new_p.append(new_r)
            
            new_tc.append(new_p)
            
            # Insert the new cell
            tr.insert(insert_idx + i, new_tc)
    
    return table


def main():
    input_file = "Content-Schedule-Template-v2.docx"
    output_file = "Content-Schedule-Template-v3.docx"
    
    print(f"📄 Opening: {input_file}")
    doc = Document(input_file)
    
    # Find the posts schedule table
    posts_table = None
    posts_table_idx = None
    for i, table in enumerate(doc.tables):
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "title" in headers and "type" in headers:
            posts_table = table
            posts_table_idx = i
            print(f"   Found posts schedule table (Table {i+1})")
            print(f"   Current columns ({len(headers)}): {[cell.text.strip() for cell in table.rows[0].cells]}")
            break
    
    if not posts_table:
        print("❌ Could not find posts schedule table!")
        return
    
    # Check if columns already exist
    headers = [cell.text.strip().lower() for cell in posts_table.rows[0].cells]
    if "photo links" in headers or "caption" in headers:
        print("ℹ️  Photo Links/Caption columns already exist. No changes needed.")
        return
    
    # Find position to insert (after "Time" column)
    time_idx = -1
    for idx, h in enumerate(headers):
        if "time" in h:
            time_idx = idx
            break
    
    if time_idx == -1:
        # Insert after column 2 (assuming Title, Date, Time order)
        time_idx = 2
    
    print(f"\n🔧 Adding columns after position {time_idx} ('{headers[time_idx]}')")
    
    # Add the new columns
    new_columns = ["Photo Links", "Caption"]
    add_columns_to_table(posts_table, time_idx, new_columns)
    
    # Verify the changes
    new_headers = [cell.text.strip() for cell in posts_table.rows[0].cells]
    print(f"   New columns ({len(new_headers)}): {new_headers}")
    
    # Save the modified document
    doc.save(output_file)
    
    # Remove draft post block tables (Tables 3-7 are typically post blocks)
    print(f"\n🗑️  Removing draft post block tables...")
    tables_to_remove = []
    for i, table in enumerate(doc.tables):
        # Check if this is a post detail block (2-column small tables with Title/Caption)
        if len(table.rows) >= 3 and len(table.rows) <= 6:
            first_col_labels = [row.cells[0].text.strip().lower() if len(row.cells) >= 2 else "" for row in table.rows]
            if "title" in first_col_labels and ("caption" in first_col_labels or "description" in first_col_labels):
                # Check if it's a template placeholder (title starts with [) - keep those
                title_value = ""
                for row in table.rows:
                    if len(row.cells) >= 2 and row.cells[0].text.strip().lower() == "title":
                        title_value = row.cells[1].text.strip()
                        break
                if title_value.startswith("["):
                    continue  # Keep placeholder templates
                tables_to_remove.append(table)
    
    # Remove the identified tables
    for table in tables_to_remove:
        tbl = table._tbl
        parent = tbl.getparent()
        parent.remove(tbl)
    
    print(f"   Removed {len(tables_to_remove)} draft post block tables")
    
    # Save the modified document
    doc.save(output_file)
    
    print(f"\n✅ Template saved: {output_file}")
    print(f"\n📋 Next steps:")
    print(f"   1. Open {output_file} in Word")
    print(f"   2. Adjust column widths as needed (Photo Links and Caption may need to be wider)")
    print(f"   3. Fill in Photo Links and Captions directly in the schedule table")

if __name__ == "__main__":
    main()
